import os
import json
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog  # para solicitar senha
from tkinter.font import Font
from datetime import datetime
import matplotlib.pyplot as plt
from collections import defaultdict
from docx import Document
from datetime import datetime
import csv
from tkinter import filedialog



from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

DATA_FILE = 'inventory.json'
HISTORY_FILE = 'history.json'
TERMS_DIR = 'terms'

if not os.path.exists(TERMS_DIR):
    os.makedirs(TERMS_DIR)

# Usuário "Mãe" e senha
ADMIN_USER = "mãe"
ADMIN_PASS = "Revalle@123"

# Opções fixas para Centro de Custo e Revendas
CENTER_COST_OPTIONS = [
    "101 - Puxada",
    "202 - Armazém",
    "301 - Administrativo",
    "401 - Vendas",
    "501 - Entrega",
    "601 - CSC"
]

REVENDAS_OPTIONS = [
    "Revalle Juazeiro",
    "Revalle Bonfim",
    "Revalle Pernambuco",
    "Revalle Nordeste",
    "Revalle Paulo Afonso",
    "Revalle Alagoinhas",
    "Revalle Serrinha"
]

TERMO_MODELOS = {
    "Revalle Juazeiro": "modelos/termo_juazeiro.docx",
    "Revalle Bonfim": "modelos/termo_bonfim.docx",
    "Revalle Pernambuco": "modelos/termo_pernambuco.docx",
    "Revalle Nordeste": "modelos/termo_nordeste.docx",
    "Revalle Paulo Afonso": "modelos/termo_pauloafonso.docx",
    "Revalle Alagoinhas": "modelos/termo_alagoinhas.docx",
    "Revalle Serrinha": "modelos/termo_serrinha.docx",
}


#FORMATAR DATA E CPF
    
def format_cpf(cpf) -> str:
    if not cpf:
        return "-"  # ou "" se preferir vazio
    cpf = str(cpf)  # garante que vira string
    digits = "".join(filter(str.isdigit, cpf))
    if len(digits) == 11:
        return f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"
    return cpf


def format_date(date_str: str) -> str:
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    except:
        return date_str  # se falhar, retorna original

class Inventory:
    def __init__(self, data_path=DATA_FILE, history_path=HISTORY_FILE):
        self.data_path = data_path
        self.history_path = history_path
        self.items = []
        self.history = []
        self.load()
        self.load_history()

    def load(self):
        if os.path.exists(self.data_path):
            with open(self.data_path, 'r', encoding='utf-8') as f:
                self.items = json.load(f)
        else:
            self.items = []
            self.save()

    def save(self):
        with open(self.data_path, 'w', encoding='utf-8') as f:
            json.dump(self.items, f, ensure_ascii=False, indent=2)

    def load_history(self):
        if os.path.exists(self.history_path):
            with open(self.history_path, 'r', encoding='utf-8') as f:
                self.history = json.load(f)
        else:
            self.history = []
            self.save_history()

    def save_history(self):
        with open(self.history_path, 'w', encoding='utf-8') as f:
            json.dump(self.history, f, ensure_ascii=False, indent=2)

    def next_id(self):
        return 1 if not self.items else max(p['id'] for p in self.items) + 1

    def add_item(self, tipo, brand, model, identificador, **kwargs):
        # Impede duplicado
        if any(p['identificador'] == identificador for p in self.items if identificador):
            return False, f"{tipo} com esse identificador já cadastrado."

        date_cadastro = datetime.now().strftime("%Y-%m-%d")

        item = {
            'id': self.next_id(),
            'tipo': tipo,
            'brand': brand,
            'model': model,
            'identificador': identificador,
            'status': 'Disponível',
            'assigned_to': None,
            'date_registered': date_cadastro,
            'date_issued': None,
            'date_returned': None,
            'center_cost': None,
            'cargo': None,
            'revenda': kwargs.get('revenda'),
            'dominio': kwargs.get('dominio'),
            'host': kwargs.get('host'),
            'endereco_fisico': kwargs.get('endereco_fisico'),
            'cpu': kwargs.get('cpu'),
            'ram': kwargs.get('ram'),
            'storage': kwargs.get('storage'),
            'sistema': kwargs.get('sistema'),
            'licenca': kwargs.get('licenca'),
            'anydesk': kwargs.get('anydesk'),
            'setor': kwargs.get('setor'),
            'ip': kwargs.get('ip'),
            'mac': kwargs.get('mac'),
        }

        self.items.append(item)
        self.save()
        return True, item


    def find(self, pid):
        return next((p for p in self.items if p['id'] == pid), None)

    def issue(self, pid, user, cpf, center_cost, cargo, revenda, date_emprestimo_str):
        p = self.find(pid)
        if not p:
            return False, "ID não encontrado."
        if p['status'] != 'Disponível':
            return False, "Aparelho já emprestado."

        try:
            dt_issue = datetime.strptime(date_emprestimo_str, "%d/%m/%Y")
        except:
            return False, "Data de empréstimo inválida. Use o formato dd/mm/aaaa."

        dt_cadastro = datetime.strptime(p['date_registered'], "%Y-%m-%d")
        if dt_issue < dt_cadastro:
            return False, "Data de empréstimo não pode ser anterior à data de cadastro."
        if p['date_returned'] and dt_issue < datetime.strptime(p['date_returned'], "%Y-%m-%d"):
            return False, "Data de empréstimo não pode ser anterior à data de devolução."

        # Atualiza o item
        p['status'] = 'Indisponível'
        p['assigned_to'] = user
        p['cpf'] = cpf
        p['date_issued'] = dt_issue.strftime("%Y-%m-%d")
        p['date_returned'] = None
        p['center_cost'] = center_cost
        p['cargo'] = cargo
        p['revenda'] = revenda

        # Log completo
        log = {
            "id": pid,
            "tipo": p.get("tipo", ""),
            "marca": p.get("brand", ""),
            "modelo": p.get("model", ""),
            "identificador": p.get("identificador", ""),
            "usuario": user,
            "cpf": cpf,
            "cargo": cargo,
            "center_cost": center_cost,
            "revenda": revenda,
            "data_emprestimo": dt_issue.strftime("%Y-%m-%d"),
            "data_devolucao": ""
        }
        self.history.append(log)

        self.save()
        self.save_history()
        return True, f"Aparelho {pid} emprestado para {user}."


    def ret(self, pid, date_devolucao_str):
        p = self.find(pid)
        if not p:
            return False, "ID não encontrado."
        if p['status'] != 'Indisponível':
            return False, "Aparelho não está emprestado."

        try:
            dt_return = datetime.strptime(date_devolucao_str, "%d/%m/%Y")
        except:
            return False, "Data de devolução inválida. Use o formato dd/mm/aaaa."

        dt_issue = datetime.strptime(p['date_issued'], "%Y-%m-%d")
        if dt_return < dt_issue:
            return False, "Data de devolução não pode ser anterior à data de empréstimo."

        # Atualiza item
        p['date_returned'] = dt_return.strftime("%Y-%m-%d")
        p['status'] = 'Disponível'
        p['assigned_to'] = None
        p['cpf'] = None

        # Atualiza histórico (preenche devolução no último empréstimo aberto)
        for h in reversed(self.history):
            if h["id"] == pid and not h.get("data_devolucao"):
                h["data_devolucao"] = dt_return.strftime("%Y-%m-%d")
                break

        self.save()
        self.save_history()
        return True, f"Aparelho {pid} devolvido."



    def remove(self, pid, user):
        p = self.find(pid)
        if not p:
            return False, "ID não encontrado."
        if p['status'] != 'Disponível':
            return False, "Não é possível remover um produto emprestado."
        self.items.remove(p)
        self.save()
        return True, f"Aparelho {pid} removido."

    def search(self, term):
        term = term.lower()
        return [
            p for p in self.items
            if term in p['brand'].lower()
            or term in p['model'].lower()
            or term in p['identificador']
        ]

    def generate_monthly_report(self, year, month):
        rows = []
        for h in self.history:
            try:
                if h.get("data_emprestimo"):
                    d1 = datetime.strptime(h["data_emprestimo"], "%Y-%m-%d")
                    if d1.year == year and d1.month == month:
                        rows.append({
                            "item_id": h["id"],
                            "operation": "Empréstimo",
                            "date": h["data_emprestimo"],
                            "user": h.get("usuario", "-"),
                            "cpf": h.get("cpf", "-"),
                            "center_cost": h.get("center_cost", "-"),
                            "cargo": h.get("cargo", "-"),
                            "revenda": h.get("revenda", "-"),
                        })
                if h.get("data_devolucao"):
                    d2 = datetime.strptime(h["data_devolucao"], "%Y-%m-%d")
                    if d2.year == year and d2.month == month:
                        rows.append({
                            "item_id": h["id"],
                            "operation": "Devolução",
                            "date": h["data_devolucao"],
                            "user": h.get("usuario", "-"),
                            "cpf": h.get("cpf", "-"),
                            "center_cost": h.get("center_cost", "-"),
                            "cargo": h.get("cargo", "-"),
                            "revenda": h.get("revenda", "-"),
                        })
            except:
                continue
        rows.sort(key=lambda r: (r["date"], r["operation"]))
        return rows


    def generate_term(self, item_id, user):
        item = self.find(item_id)
        if not item:
            return False, "Equipamento não encontrado."
        if item['status'] != 'Indisponível' or item['assigned_to'] != user:
            return False, "Equipamento não está emprestado para este usuário."

        revenda = item.get("revenda")
        modelo_path = TERMO_MODELOS.get(revenda)
        if not modelo_path or not os.path.exists(modelo_path):
            return False, f"Modelo de termo não encontrado para {revenda}."

        safe_user_name = user.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saida_path = os.path.join(TERMS_DIR, f"termo_{item_id}_{safe_user_name}_{revenda}_{timestamp}.docx")

        doc = Document(modelo_path)

        substituicoes = {
            "{{nome}}": user,
            "{{data_hoje}}": datetime.now().strftime("%d/%m/%Y"),
        }

        for key, value in item.items():
            placeholder = f"{{{{{key}}}}}"
            str_value = str(value) if value not in [None, ''] else ""
            substituicoes[placeholder] = str_value

        substituicoes["{{cpf}}"] = format_cpf(item.get("cpf", ""))
        substituicoes["{{data_cadastro}}"] = format_date(item.get("date_registered", ""))
        substituicoes["{{data_emprestimo}}"] = format_date(item.get("date_issued", ""))
        substituicoes["{{marca}}"] = item.get("brand", "")
        substituicoes["{{modelo}}"] = item.get("model", "")
        substituicoes["{{tipo}}"] = item.get("tipo", "")
        substituicoes["{{identificador}}"] = item.get("identificador", "")
        

        for p in doc.paragraphs:
            for chave, valor in substituicoes.items():
                if chave in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if chave in inline[i].text:
                            text = inline[i].text.replace(chave, str(valor))
                            inline[i].text = text

        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        for chave, valor in substituicoes.items():
                            if chave in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if chave in inline[i].text:
                                        text = inline[i].text.replace(chave, str(valor))
                                        inline[i].text = text
        doc.save(saida_path)
        return True, saida_path

    def recalc_item_status(self, item_id):
        """
        Recalcula o status do aparelho com base nos lançamentos (issue e return) presentes no histórico.
        """
        item = self.find(item_id)
        if not item:
            return
        registros = [h for h in self.history if h.get("id") == item_id]
        if not registros:
            item['status'] = 'Disponível'
            item['assigned_to'] = None
            item['cpf'] = None
            item['date_issued'] = None
            item['center_cost'] = None
            item['cargo'] = None
            item['revenda'] = None
            self.save()
            return

        # último empréstimo aberto (sem devolução). Se houver, o item está indisponível.
        abertos = [h for h in registros if not h.get("data_devolucao")]
        if abertos:
            # pega o mais recente por data_emprestimo
            h = max(abertos, key=lambda x: x.get("data_emprestimo") or "")
            item['status'] = 'Indisponível'
            item['assigned_to'] = h.get('usuario')
            item['cpf'] = h.get('cpf')
            item['date_issued'] = h.get('data_emprestimo')
            item['center_cost'] = h.get('center_cost')
            item['cargo'] = h.get('cargo')
            item['revenda'] = h.get('revenda')
        else:
            item['status'] = 'Disponível'
            item['assigned_to'] = None
            item['cpf'] = None
            item['date_issued'] = None
            item['center_cost'] = None
            item['cargo'] = None
            item['revenda'] = None
        self.save()

    def get_issue_return_counts(self):
        """
        Retorna dois dicionários com a quantidade de operações de "issue" e "return" agrupadas por YYYY-MM.
        """
        issues = defaultdict(int)
        returns = defaultdict(int)
        for h in self.history:
            if h.get("data_emprestimo"):
                key = h["data_emprestimo"][:7]
                issues[key] += 1
            if h.get("data_devolucao"):
                key = h["data_devolucao"][:7]
                returns[key] += 1
        return issues, returns

    def get_registration_counts(self):
        """
        Retorna um dicionário com a quantidade de cadastros agrupados por YYYY-MM.
        """
        regs = defaultdict(int)
        for item in self.items:
            key = item.get('date_registered', "")[:7]
            if key:
                regs[key] += 1
        return regs

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Gestão de Estoque de Celulares")
        self.geometry("1000x750")
        self.inv = Inventory()
        self.create_widgets()

    def create_widgets(self):
        notebook = ttk.Notebook(self)
        notebook.pack(fill='both', expand=True)

        self.tab_stock   = ttk.Frame(notebook)
        self.tab_add     = ttk.Frame(notebook)
        self.tab_issue   = ttk.Frame(notebook)
        self.tab_return  = ttk.Frame(notebook)
        self.tab_remove  = ttk.Frame(notebook)
        self.tab_report  = ttk.Frame(notebook)
        self.tab_history = ttk.Frame(notebook)
        self.tab_terms   = ttk.Frame(notebook)
        self.tab_graphs  = ttk.Frame(notebook)

        notebook.add(self.tab_stock,  text="Estoque")
        notebook.add(self.tab_add,    text="Cadastrar")
        notebook.add(self.tab_issue,  text="Emprestar")
        notebook.add(self.tab_return, text="Devolver")
        notebook.add(self.tab_remove, text="Remover")
        notebook.add(self.tab_report, text="Relatório")
        notebook.add(self.tab_history,text="Histórico")
        notebook.add(self.tab_terms,  text="Termos")
        notebook.add(self.tab_graphs, text="Gráficos")

        self.build_stock_tab()
        self.build_add_tab()
        self.build_issue_tab()
        self.build_return_tab()
        self.build_remove_tab()
        self.build_report_tab()
        self.build_history_tab()
        self.build_terms_tab()
        self.build_graph_tab()

    def autosize_treeview(self, tree):
        font = Font()
        for col in tree["columns"]:
            max_width = font.measure(col) + 10
            for item in tree.get_children():
                cell_val = tree.set(item, col)
                cell_width = font.measure(cell_val) + 10
                if cell_width > max_width:
                    max_width = cell_width
            tree.column(col, width=max_width)

    def build_stock_tab(self):
        tab = self.tab_stock  # usa a aba estoque já criada no create_widgets

        # --- Estilo do Treeview ---
        style = ttk.Style()
        style.configure("Treeview", rowheight=28, font=("Segoe UI", 10))
        style.configure("Treeview.Heading", font=("Segoe UI", 11, "bold"))

        # --- Filtros ---
        frm_filters = ttk.Frame(tab)
        frm_filters.pack(fill="x", padx=10, pady=5)

        ttk.Label(frm_filters, text="Status:").grid(row=0, column=0, padx=5, pady=2, sticky="e")
        self.cb_status_filter = ttk.Combobox(frm_filters, values=["", "Disponível", "Indisponível"], state="readonly", width=15)
        self.cb_status_filter.grid(row=0, column=1, padx=5, pady=2)

        ttk.Label(frm_filters, text="Tipo:").grid(row=0, column=2, padx=5, pady=2, sticky="e")
        self.cb_tipo_filter = ttk.Combobox(frm_filters, values=["", "Celular", "Notebook", "Desktop", "Impressora", "Tablet"], state="readonly", width=15)
        self.cb_tipo_filter.grid(row=0, column=3, padx=5, pady=2)

        ttk.Label(frm_filters, text="Revenda:").grid(row=0, column=4, padx=5, pady=2, sticky="e")
        self.cb_revenda_filter = ttk.Combobox(frm_filters, values=[""] + REVENDAS_OPTIONS, state="readonly", width=20)
        self.cb_revenda_filter.grid(row=0, column=5, padx=5, pady=2)

        ttk.Label(frm_filters, text="Buscar:").grid(row=0, column=6, padx=5, pady=2, sticky="e")
        self.e_search_stock = ttk.Entry(frm_filters, width=20)
        self.e_search_stock.grid(row=0, column=7, padx=5, pady=2)


        ttk.Button(frm_filters, text="Aplicar", command=self.update_stock).grid(row=0, column=8, padx=10, pady=2)
        ttk.Button(frm_filters, text="Limpar", command=self.clear_filters).grid(row=0, column=9, padx=5, pady=2)
        ttk.Button(frm_filters, text="Exportar CSV", command=lambda: self.exportar_csv(self.tree_stock, "Exportar Estoque", "estoque")).grid(row=0, column=10, padx=5, pady=2)


        
        # --- Treeview ---
        cols = [
            "ID",
            "Revenda", "Tipo", "Marca", "Modelo", "Status", "Usuário", "CPF",
            "Identificador", "Domínio", "Host", "Endereço Físico", "CPU",
            "RAM", "Storage", "Sistema", "Licença", "AnyDesk",
            "Setor", "IP", "MAC", "Data Cadastro"
        ]

        self.tree_stock = ttk.Treeview(tab, columns=cols, show="headings", height=18)
        self.tree_stock.pack(fill="both", expand=True, padx=10, pady=5)

        col_widths = {
            "ID": 30,
            "Revenda": 120, "Tipo": 100, "Marca": 100, "Modelo": 120, "Status": 100,
            "Usuário": 140, "CPF" : 100, "Identificador": 140, "Domínio": 100, "Host": 140,
            "Endereço Físico": 160, "CPU": 120, "RAM": 100, "Storage": 100,
            "Sistema": 140, "Licença": 140, "AnyDesk": 120, "Setor": 100,
            "IP": 120, "MAC": 140, "Data Cadastro": 140,
        }

        for col in cols:
            self.tree_stock.heading(col, text=col, command=lambda c=col: self.treeview_sort_column(self.tree_stock, c, False))
            self.tree_stock.column(col, width=col_widths.get(col, 100), anchor="w", stretch=False)

        vsb = ttk.Scrollbar(tab, orient="vertical", command=self.tree_stock.yview)
        hsb = ttk.Scrollbar(tab, orient="horizontal", command=self.tree_stock.xview)
        self.tree_stock.configure(yscroll=vsb.set, xscroll=hsb.set)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")

        self.update_stock()


    def clear_filters(self):
        self.cb_status_filter.set("")
        self.cb_tipo_filter.set("")
        self.cb_revenda_filter.set("")
        self.e_search_stock.delete(0, "end")  # <- corrigido (era e_search.stock)
        self.update_stock()



    def build_add_tab(self):
        frm = ttk.Frame(self.tab_add, padding=10)
        frm.pack(fill="both", expand=True)

        # --- Campo Tipo ---
        ttk.Label(frm, text="Tipo de Equipamento:").grid(row=0, column=0, sticky="e")
        self.cb_tipo = ttk.Combobox(frm, values=["Celular", "Notebook", "Desktop", "Impressora", "Tablet"], state="readonly")
        self.cb_tipo.grid(row=0, column=1)
        self.cb_tipo.bind("<<ComboboxSelected>>", self.on_tipo_selected)

        # Frame onde os campos adicionais vão aparecer
        self.frm_dynamic = ttk.Frame(frm, padding=10)
        self.frm_dynamic.grid(row=1, column=0, columnspan=2, sticky="nsew")

        # Mensagem de feedback
        self.lbl_add = ttk.Label(frm, text="", foreground="red")
        self.lbl_add.grid(row=99, column=0, columnspan=2)

        # Botão cadastrar
        ttk.Button(frm, text="Cadastrar", command=self.cmd_add).grid(row=100, column=0, columnspan=2, pady=5)
    
    
    
    def on_tipo_selected(self, event=None):
        # Limpa os widgets anteriores
        for widget in self.frm_dynamic.winfo_children():
            widget.destroy()

        tipo = self.cb_tipo.get()

        # --- Campos comuns ---
        ttk.Label(self.frm_dynamic, text="Marca:").grid(row=1, column=0, sticky="e")
        self.e_brand = ttk.Entry(self.frm_dynamic)
        self.e_brand.grid(row=1, column=1)

        ttk.Label(self.frm_dynamic, text="Revenda:").grid(row=2, column=0, sticky="e")
        self.cb_revenda_add = ttk.Combobox(
            self.frm_dynamic, values=REVENDAS_OPTIONS, state="readonly"
        )
        self.cb_revenda_add.grid(row=2, column=1)

        # --- Campos específicos por tipo ---
        if tipo == "Celular":
            ttk.Label(self.frm_dynamic, text="Modelo:").grid(row=3, column=0, sticky="e")
            self.e_model = ttk.Entry(self.frm_dynamic)
            self.e_model.grid(row=3, column=1)
            
            ttk.Label(self.frm_dynamic, text="IMEI:").grid(row=4, column=0, sticky="e")
            self.e_identificador = ttk.Entry(self.frm_dynamic)
            self.e_identificador.grid(row=4, column=1)

        elif tipo in ["Notebook", "Desktop"]:
            ttk.Label(self.frm_dynamic, text="Domínio:").grid(row=3, column=0, sticky="e")
            self.cb_dominio = ttk.Combobox(
                self.frm_dynamic, values=["Sim", "Não"], state="readonly"
            )
            self.cb_dominio.grid(row=3, column=1)

            ttk.Label(self.frm_dynamic, text="Host:").grid(row=4, column=0, sticky="e")
            self.e_host = ttk.Entry(self.frm_dynamic)
            self.e_host.grid(row=4, column=1)

            ttk.Label(self.frm_dynamic, text="Endereço Físico:").grid(row=5, column=0, sticky="e")
            self.e_endereco_fisico = ttk.Entry(self.frm_dynamic)
            self.e_endereco_fisico.grid(row=5, column=1)

            ttk.Label(self.frm_dynamic, text="Armazenamento (GB):").grid(row=6, column=0, sticky="e")
            self.e_storage = ttk.Entry(self.frm_dynamic)
            self.e_storage.grid(row=6, column=1)

            ttk.Label(self.frm_dynamic, text="Sistema Operacional:").grid(row=7, column=0, sticky="e")
            self.e_sistema = ttk.Entry(self.frm_dynamic)
            self.e_sistema.grid(row=7, column=1)

            ttk.Label(self.frm_dynamic, text="Processador:").grid(row=8, column=0, sticky="e")
            self.e_cpu = ttk.Entry(self.frm_dynamic)
            self.e_cpu.grid(row=8, column=1)

            ttk.Label(self.frm_dynamic, text="Memória RAM:").grid(row=9, column=0, sticky="e")
            self.e_ram = ttk.Entry(self.frm_dynamic)
            self.e_ram.grid(row=9, column=1)

            ttk.Label(self.frm_dynamic, text="Licença Windows:").grid(row=10, column=0, sticky="e")
            self.e_licenca = ttk.Entry(self.frm_dynamic)
            self.e_licenca.grid(row=10, column=1)

            ttk.Label(self.frm_dynamic, text="AnyDesk:").grid(row=11, column=0, sticky="e")
            self.e_anydesk = ttk.Entry(self.frm_dynamic)
            self.e_anydesk.grid(row=11, column=1)

        elif tipo == "Impressora":
            ttk.Label(self.frm_dynamic, text="Setor:").grid(row=3, column=0, sticky="e")
            self.e_setor = ttk.Entry(self.frm_dynamic)
            self.e_setor.grid(row=3, column=1)

            ttk.Label(self.frm_dynamic, text="IP:").grid(row=4, column=0, sticky="e")
            self.e_ip = ttk.Entry(self.frm_dynamic)
            self.e_ip.grid(row=4, column=1)

            ttk.Label(self.frm_dynamic, text="MAC:").grid(row=5, column=0, sticky="e")
            self.e_mac = ttk.Entry(self.frm_dynamic)
            self.e_mac.grid(row=5, column=1)

        elif tipo == "Tablet":
            ttk.Label(self.frm_dynamic, text="Nº de Série:").grid(row=3, column=0, sticky="e")
            self.e_identificador = ttk.Entry(self.frm_dynamic)
            self.e_identificador.grid(row=3, column=1)

            ttk.Label(self.frm_dynamic, text="Armazenamento (GB):").grid(row=4, column=0, sticky="e")
            self.e_storage = ttk.Entry(self.frm_dynamic)
            self.e_storage.grid(row=4, column=1)




    def build_issue_tab(self):
        frm = ttk.Frame(self.tab_issue, padding=10)
        frm.pack()
        ttk.Label(frm, text="Selecione aparelho:").grid(row=0, column=0, sticky='e')
        self.cb_issue       = ttk.Combobox(frm, state='readonly', width=30)
        self.cb_issue.grid(row=0, column=1)
        
        ttk.Label(frm, text="Funcionário:").grid(row=1, column=0, sticky='e')
        self.e_issue_user = ttk.Entry(frm); self.e_issue_user.grid(row=1, column=1)

        ttk.Label(frm, text="CPF:").grid(row=2, column=0, sticky='e')
        self.e_issue_cpf = ttk.Entry(frm); self.e_issue_cpf.grid(row=2, column=1)

        ttk.Label(frm, text="Centro de Custo:").grid(row=3, column=0, sticky='e')
        self.cb_center      = ttk.Combobox(frm, state='readonly', values=CENTER_COST_OPTIONS, width=28)
        self.cb_center.grid(row=3, column=1)
        
        ttk.Label(frm, text="Cargo:").grid(row=4, column=0, sticky='e')
        self.e_cargo        = ttk.Entry(frm); self.e_cargo.grid(row=4, column=1)
        
        ttk.Label(frm, text="Revenda:").grid(row=5, column=0, sticky='e')
        self.cb_revenda     = ttk.Combobox(frm, state='readonly', values=REVENDAS_OPTIONS, width=28)
        self.cb_revenda.grid(row=5, column=1)
        
        ttk.Label(frm, text="Data Empréstimo (dd/mm/aaaa):").grid(row=6, column=0, sticky='e')
        self.e_date_issue   = ttk.Entry(frm); self.e_date_issue.grid(row=6, column=1)
        self.lbl_issue      = ttk.Label(frm, text="", foreground='red')
        self.lbl_issue.grid(row=7, column=0, columnspan=2)
        
        ttk.Button(frm, text="Emprestar", command=self.cmd_issue).grid(row=8, column=0, columnspan=2, pady=5)
        self.update_issue_cb()
        
        self.e_issue_cpf.bind("<KeyRelease>", self.on_cpf_entry)
        self.e_date_issue.bind("<KeyRelease>", lambda e: self.on_date_entry(e, self.e_date_issue))


    def build_return_tab(self):
        frm = ttk.Frame(self.tab_return, padding=10)
        frm.pack()

        ttk.Label(frm, text="Selecione aparelho:").grid(row=0, column=0, sticky='e')
        self.cb_return = ttk.Combobox(frm, state='readonly', width=30)
        self.cb_return.grid(row=0, column=1)

        ttk.Label(frm, text="Data Devolução (dd/mm/aaaa):").grid(row=1, column=0, sticky='e')
        self.e_date_return = ttk.Entry(frm)
        self.e_date_return.grid(row=1, column=1)

        # 🔹 Aqui entra o bind da máscara de data
        self.e_date_return.bind("<KeyRelease>", lambda e: self.on_date_entry(e, self.e_date_return))

        self.lbl_ret = ttk.Label(frm, text="", foreground='red')
        self.lbl_ret.grid(row=2, column=0, columnspan=2)

        ttk.Button(frm, text="Devolver", command=self.cmd_return).grid(row=3, column=0, columnspan=2, pady=5)

        self.update_return_cb()


    def build_remove_tab(self):
        frm = ttk.Frame(self.tab_remove, padding=10)
        frm.pack()
        ttk.Label(frm, text="Selecione aparelho:").grid(row=0, column=0, sticky='e')
        self.cb_remove      = ttk.Combobox(frm, state='readonly', width=30)
        self.cb_remove.grid(row=0, column=1)
        self.lbl_rem        = ttk.Label(frm, text="", foreground='red')
        self.lbl_rem.grid(row=1, column=0, columnspan=2)
        ttk.Button(frm, text="Remover", command=self.cmd_remove).grid(row=2, column=0, columnspan=2, pady=5)
        self.update_remove_cb()


    def build_report_tab(self):
        frm = ttk.Frame(self.tab_report, padding=10)
        frm.pack()
        top_frm = ttk.Frame(frm); top_frm.pack(fill='x')
        ttk.Label(top_frm, text="Ano:").grid(row=0, column=0, sticky='e')
        self.e_report_year  = ttk.Entry(top_frm, width=6)
        self.e_report_year.insert(0, datetime.now().year)
        self.e_report_year.grid(row=0, column=1)
        ttk.Label(top_frm, text="Mês:").grid(row=0, column=2, sticky='e', padx=5)
        self.cb_report_month = ttk.Combobox(top_frm, values=list(range(1, 13)), width=4)
        self.cb_report_month.set(datetime.now().month)
        self.cb_report_month.grid(row=0, column=3)
        ttk.Button(top_frm, text="Gerar Relatório", command=self.cmd_generate_report).grid(row=0, column=4, padx=10)
        ttk.Button(top_frm, text="Estornar Lançamento", command=self.cmd_delete_report_entry).grid(row=0, column=5, padx=10)
        ttk.Button(top_frm, text="Exportar CSV", command=lambda: self.exportar_csv(self.tree_report, "Exportar Relatório", "relatório")).grid(row=0, column=6, padx=5, pady=2)

        
        cols = (
        "ID", "Tipo", "Marca", "Modelo", "Identificador", 
        "Usuário", "CPF", "Operação", "Data", 
        "Centro de Custo", "Cargo", "Revenda"
        )

        self.tree_report = ttk.Treeview(self.tab_report, columns=cols, show='headings')
        larguras = {
        "ID": 60,
        "Tipo": 100,
        "Marca": 120,
        "Modelo": 140,
        "Identificador": 160,
        "Usuário": 200,
        "CPF": 120,
        "Operação": 110,
        "Data": 100,
        "Centro de Custo": 160,
        "Cargo": 140,
        "Revenda": 140,
        }

        for c in cols:
            self.tree_report.heading(c, text=c)
            self.tree_report.column(c, width=larguras.get(c, 120), anchor='w', stretch=False)
        self.tree_report.pack(fill='both', expand=True, padx=5, pady=5)
        
        

    def exportar_csv(self, tree, titulo="Exportar", nome="dados"):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv")]
        )
        if not file_path:
            return

        with open(file_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)

            headers = [tree.heading(c)["text"] for c in tree["columns"]]
            writer.writerow(headers)

            for row_id in tree.get_children():
                values = tree.item(row_id)["values"]
                writer.writerow(values)

        messagebox.showinfo(titulo, f"{nome.capitalize()} exportado para:\n{file_path}")

        
        
    #Ordenação Cabeçalho
    def treeview_sort_column(self, tv, col, reverse):
        items = [(tv.set(k, col), k) for k in tv.get_children("")]
        try:
            items.sort(key=lambda t: int(t[0]) if t[0].isdigit() else t[0], reverse=reverse)
        except:
            items.sort(reverse=reverse)
        for index, (val, k) in enumerate(items):
            tv.move(k, "", index)
        tv.heading(col, command=lambda: self.treeview_sort_column(tv, col, not reverse))
        
        
    def build_history_tab(self):
        tab = self.tab_history  # usa a aba já criada em create_widgets

        top = ttk.Frame(tab)
        top.pack(fill="x", padx=10, pady=5)
        ttk.Button(top, text="Exportar CSV",
                command=lambda: self.exportar_csv(self.tree_history, "Exportar Histórico", "histórico")
                ).pack(side="right")

        cols = (
            "ID", "Tipo", "Marca", "Modelo", "Identificador", "Usuário", "CPF",
            "Cargo", "Centro de Custo", "Revenda",
            "Data Empréstimo", "Data Devolução"
        )
        self.tree_history = ttk.Treeview(tab, columns=cols, show="headings", height=18)
        widths = {
            "ID": 60, "Tipo": 100, "Marca": 120, "Modelo": 140, "Identificador": 140,
            "Usuário": 200, "CPF": 130, "Cargo": 140, "Centro de Custo": 160,
            "Revenda": 140, "Data Empréstimo": 130, "Data Devolução": 130
        }
        for c in cols:
            self.tree_history.heading(c, text=c)
            self.tree_history.column(c, width=widths.get(c, 120), anchor="w", stretch=False)

        ysb = ttk.Scrollbar(tab, orient="vertical", command=self.tree_history.yview)
        xsb = ttk.Scrollbar(tab, orient="horizontal", command=self.tree_history.xview)
        self.tree_history.configure(yscroll=ysb.set, xscroll=xsb.set)

        self.tree_history.pack(fill="both", expand=True, padx=10, pady=(0,5))
        ysb.pack(side="right", fill="y")
        xsb.pack(side="bottom", fill="x")

        # carrega dados ao abrir
        self.update_history_table()




    def build_terms_tab(self):
        frm = ttk.Frame(self.tab_terms, padding=10)
        frm.pack(fill="both", expand=True)

        # --- Filtro de busca ---
        search_frame = ttk.Frame(frm)
        search_frame.pack(fill="x", pady=5)
        ttk.Label(search_frame, text="Buscar:").grid(row=0, column=0, sticky="e")
        self.e_search_terms = ttk.Entry(search_frame, width=30)
        self.e_search_terms.grid(row=0, column=1, padx=5)
        ttk.Button(search_frame, text="Buscar", command=self.cmd_search_terms).grid(row=0, column=2, padx=5)

        # --- Tabela de termos (empréstimos ativos) ---
        cols = ("ID", "Tipo", "Marca", "Usuário", "CPF", "Data Empréstimo", "Revenda")

        container = ttk.Frame(frm)
        container.pack(fill="both", expand=True, pady=10)

        xscroll = ttk.Scrollbar(container, orient="horizontal")
        yscroll = ttk.Scrollbar(container, orient="vertical")

        self.tree_terms = ttk.Treeview(
            container,
            columns=cols,
            show="headings",
            displaycolumns=cols,
            xscrollcommand=xscroll.set,
            yscrollcommand=yscroll.set,
            height=12
        )

        self.tree_terms.column("#0", width=0, stretch=False)

        larguras = {
            "ID": 50,
            "Tipo": 100,
            "Marca": 100,
            "Usuário": 200,
            "CPF": 150,
            "Data Empréstimo": 150,
            "Revenda": 150,
        }

        for c in cols:
            self.tree_terms.heading(c, text=c, anchor="w")
            self.tree_terms.column(c, width=larguras[c], minwidth=larguras[c], anchor="w", stretch=False)

        yscroll.config(command=self.tree_terms.yview)
        xscroll.config(command=self.tree_terms.xview)
        yscroll.pack(side="right", fill="y")
        xscroll.pack(side="bottom", fill="x")
        self.tree_terms.pack(side="left", fill="both", expand=True)

        
        # depois de inserir as linhas:
        self.tree_terms.update_idletasks()
        for c in cols:
            atual = self.tree_terms.column(c, "width")
            minimo = larguras[c]
            if atual < minimo:
                self.tree_terms.column(c, width=minimo)


        # --- Botão para gerar termo ---
        self.lbl_terms = ttk.Label(frm, text="", foreground="red")
        self.lbl_terms.pack()
        ttk.Button(frm, text="Gerar Termo Selecionado", command=self.cmd_generate_term).pack(pady=5)

        # Carrega os empréstimos ativos
        self.update_terms_table()


    #MASKS CPF AND DATE  
    def on_cpf_entry(self, event):
        text = "".join(filter(str.isdigit, self.e_issue_cpf.get()))
        if len(text) > 11:
            text = text[:11]
        formatted = format_cpf(text)
        self.e_issue_cpf.delete(0, tk.END)
        self.e_issue_cpf.insert(0, formatted)
        self.e_issue_cpf.icursor(tk.END)

    def on_date_entry(self, event, entry_widget):
        text = "".join(filter(str.isdigit, entry_widget.get()))  # só números
        if len(text) > 8:
            text = text[:8]

        # monta a string formatada
        formatted = text
        if len(text) > 2:
            formatted = text[:2] + "/" + text[2:]
        if len(text) > 4:
            formatted = formatted[:5] + "/" + formatted[5:]

        # atualiza o campo sem perder o cursor no final
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, formatted)
        entry_widget.icursor(tk.END)


    def update_history_table(self):
        self.tree_history.delete(*self.tree_history.get_children())
        for h in self.inv.history:
            self.tree_history.insert("", "end", values=(
                h.get("id", ""),               # Use .get() for safety
                h.get("tipo", ""),             # Use .get() for safety
                h.get("marca", ""),            # Use .get() for safety
                h.get("modelo", ""),           # Added missing field
                h.get("identificador", ""),    # Added missing field
                h.get("usuario", ""),          # Use .get() for safety
                format_cpf(h.get("cpf", "")),
                h.get("cargo", ""),
                h.get("center_cost", ""),
                h.get("revenda", ""),
                format_date(h.get("data_emprestimo", "")),
                format_date(h.get("data_devolucao", "")) if h.get("data_devolucao") else ""
            ))

    

    def update_terms_table(self, search_text=""):
        self.tree_terms.delete(*self.tree_terms.get_children())
        for p in self.inv.items:
            if p['status'] == 'Indisponível':
                row_str = " ".join(str(v).lower() for v in p.values())
                if search_text and search_text not in row_str:
                    continue
                row = (
                    p.get('id', ''),
                    p.get('tipo', ''),
                    p.get('brand', ''),
                    p.get('assigned_to', ''),
                    format_cpf(p.get('cpf', '')),
                    format_date(p.get('date_issued', '')),
                    p.get('revenda', '')        # agora aparece a revenda
                )
                self.tree_terms.insert('', 'end', values=row)

  
                
    def cmd_search_terms(self):
        term = self.e_search_terms.get().strip().lower()
        self.update_terms_table(term)



    def build_graph_tab(self):
        frm = ttk.Frame(self.tab_graphs, padding=10)
        frm.pack(fill='both', expand=True)

        # Seleção de Ano e Mês
        ctrl = ttk.Frame(frm)
        ctrl.pack(pady=5)
        ttk.Label(ctrl, text="Ano:").pack(side='left')
        self.e_graph_year = ttk.Entry(ctrl, width=6)
        self.e_graph_year.insert(0, datetime.now().year)
        self.e_graph_year.pack(side='left', padx=(5, 15))
        ttk.Label(ctrl, text="Mês:").pack(side='left')
        self.cb_graph_month = ttk.Combobox(ctrl,
            values=list(range(1, 13)), width=4, state='readonly')
        self.cb_graph_month.set(datetime.now().month)
        self.cb_graph_month.pack(side='left', padx=5)

        # Botões de Gráficos
        lbl = ttk.Label(frm, text="Gráficos:")
        lbl.pack(pady=(20, 5))
        btns = ttk.Frame(frm)
        btns.pack()
        ttk.Button(btns,
            text="Empréstimos x Devoluções (Mensal)",
            command=self.graph_issue_return
        ).pack(side='left', padx=10)
        ttk.Button(btns,
            text="Cadastros (Mensal)",
            command=self.graph_registration
        ).pack(side='left', padx=10)

    # --- Atualizações de UI ---
    def update_stock(self):
        self.tree_stock.delete(*self.tree_stock.get_children())

        status_filter = self.cb_status_filter.get()
        tipo_filter = self.cb_tipo_filter.get()
        revenda_filter = self.cb_revenda_filter.get()
        search_text = self.e_search_stock.get().lower().strip()


        for p in self.inv.items:
            if status_filter and p.get('status') != status_filter:
                continue
            if tipo_filter and p.get('tipo') != tipo_filter:
                continue
            if revenda_filter and p.get('revenda') != revenda_filter:
                continue
            if search_text:
                row_str = " ".join(str(v).lower() for v in p.values())
                if search_text not in row_str:
                    continue

            row = (
                p.get('id', ''),
                p.get('revenda', ''),
                p.get('tipo', ''),
                p.get('brand', ''),
                p.get('model', ''),
                p.get('status', ''),
                p.get('assigned_to', ''),
                format_cpf(p.get('cpf', '')),
                p.get('identificador', ''),
                p.get('dominio', ''),
                p.get('host', ''),
                p.get('endereco_fisico', ''),
                p.get('cpu', ''),
                p.get('ram', ''),
                p.get('storage', ''),
                p.get('sistema', ''),
                p.get('licenca', ''),
                p.get('anydesk', ''),
                p.get('setor', ''),
                p.get('ip', ''),
                p.get('mac', ''),
                format_date(p.get('date_registered', '')),
            )
            tag = "disp" if p.get('status') == "Disponível" else "indisp"
            self.tree_stock.insert('', 'end', values=row, tags=(tag,))

        self.tree_stock.tag_configure("disp", background="#6ef38d")   # verde claro
        self.tree_stock.tag_configure("indisp", background="#f56d78") # vermelho claro


    def update_issue_cb(self):
        lst = [
            f"{p['id']} - {p['brand']} {p['model']}"
            for p in self.inv.items if p['status'] == 'Disponível'
        ]
        self.cb_issue['values'] = lst
        self.cb_issue.set('')

    def update_return_cb(self):
        lst = [
            f"{p['id']} - {p['brand']} {p['model']} ({p['assigned_to']})"
            for p in self.inv.items if p['status'] == 'Indisponível'
        ]
        self.cb_return['values'] = lst
        self.cb_return.set('')

    def update_remove_cb(self):
        lst = [f"{p['id']} - {p['brand']} {p['model']}" for p in self.inv.items]
        self.cb_remove['values'] = lst
        self.cb_remove.set('')

    def update_terms_cb(self):
        # se a combobox ainda não foi criada, não faz nada
        if not hasattr(self, "cb_terms"):
            return

        lst = []
        for p in self.inv.items:  # <- corrigido (era products)
            if p.get("assigned_to"):
                lst.append(f'{p["id"]} - {p["brand"]} ({p["assigned_to"]})')

        self.cb_terms["values"] = lst




    def safe_get(self, attr):
        """
        Retorna o conteúdo de um Entry/Combobox se ele existir e não tiver sido destruído.
        Caso contrário retorna "" (string vazia).
        """
        widget = getattr(self, attr, None)
        if widget and widget.winfo_exists():
            try:
                return widget.get().strip()
            except:
                return ""
        return ""



    # --- Comandos de Ação ---
    def cmd_add(self):
        tipo = self.cb_tipo.get().strip()
        if not tipo:
            self.lbl_add.config(text="Selecione o tipo de equipamento.", foreground="red")
            return

        b = self.safe_get("e_brand")
        m = self.safe_get("e_model")
        i = self.safe_get("e_identificador")
        revenda = self.safe_get("cb_revenda_add")
        dominio = self.safe_get("cb_dominio")
        host = self.safe_get("e_host")
        endereco_fisico = self.safe_get("e_endereco_fisico")
        cpu = self.safe_get("e_cpu")
        ram = self.safe_get("e_ram")
        storage = self.safe_get("e_storage")
        sistema = self.safe_get("e_sistema")
        licenca = self.safe_get("e_licenca")
        anydesk = self.safe_get("e_anydesk")
        setor = self.safe_get("e_setor")
        ip = self.safe_get("e_ip")
        mac = self.safe_get("e_mac")


        # --- Validação ---
        erros = []

        if not b:
            erros.append("Informe a marca.")
        if not revenda:
            erros.append("Preencha o campo Revenda.")

        if tipo in ["Celular", "Tablet"] and not i:
            erros.append("Preencha o Identificador.")

        if tipo == "Impressora":
            if not ip:
                erros.append("Preencha o IP.")
            if not setor:
                erros.append("Preencha o setor.")
            if not mac:
                erros.append("Preencha o MAC.")

        if tipo in ["Desktop", "Notebook"]:
            if not dominio:
                erros.append("Informe o domínio.")
            if not host:
                erros.append("Informe o host.")
            if not endereco_fisico:
                erros.append("Informe o endereço físico.")
            if not storage:
                erros.append("Informe o tamanho do armazenamento.")
            if not sistema:
                erros.append("Informe o sistema operacional.")
            if not cpu:
                erros.append("Informe o processador.")
            if not ram:
                erros.append("Informe a quantidade de memória RAM.")
            if not licenca:
                erros.append("Informe a licença do Windows.")
            if not anydesk:
                erros.append("Informe o código do AnyDesk.")

        if erros:
            self.lbl_add.config(text="\n".join(erros), foreground="red")
            return

        # --- Cadastro ---
        ok, res = self.inv.add_item(
            tipo, b, m, i,
            revenda=revenda,
            dominio=dominio,
            host=host,
            endereco_fisico=endereco_fisico,
            cpu=cpu,
            ram=ram,
            storage=storage,
            sistema=sistema,
            licenca=licenca,
            anydesk=anydesk,
            setor=setor,
            ip=ip,
            mac=mac,
        )
        if not ok:
            self.lbl_add.config(text=res, foreground="red")
            return

        self.lbl_add.config(text=f"{tipo} ID {res['id']} cadastrado.", foreground="green")

        # Limpar
        self.cb_tipo.set("")
        for widget in self.frm_dynamic.winfo_children():
            widget.destroy()

        # Atualizar tabelas
        self.update_stock()
        self.update_issue_cb()
        self.update_remove_cb()




    def cmd_issue(self):
        sel = self.cb_issue.get()
        user = self.e_issue_user.get().strip()
        cpf = self.e_issue_cpf.get().strip()
        cc, cargo, revenda, date_issue = (
            self.cb_center.get().strip(), self.e_cargo.get().strip(),
            self.cb_revenda.get().strip(), self.e_date_issue.get().strip()
        )
        if not (sel and user and cc and cargo and revenda and date_issue):
            self.lbl_issue.config(text="Preencha todos os campos de empréstimo.", foreground='red')
            return
        pid, _ = sel.split(' - ', 1)
        pid = int(pid)
        ok, msg = self.inv.issue(pid, user, cpf, cc, cargo, revenda, date_issue)
        self.lbl_issue.config(text=msg, foreground='green' if ok else 'red')
        if ok:
            self.e_issue_user.delete(0, 'end'); self.cb_center.set('')
            self.e_cargo.delete(0, 'end'); self.cb_revenda.set('')
            self.e_date_issue.delete(0, 'end')
            self.update_stock(); self.update_issue_cb()
            self.update_return_cb(); self.update_terms_cb()
            self.update_history_table()


    def cmd_return(self):
        sel = self.cb_return.get()
        date_return = self.e_date_return.get().strip()
        if not (sel and date_return):
            self.lbl_ret.config(text="Selecione aparelho e informe data de devolução.")
            return
        pid, _ = sel.split(' - ', 1)
        pid = int(pid)
        ok, msg = self.inv.ret(pid, date_return)
        self.lbl_ret.config(text=msg, foreground='green' if ok else 'red')
        if ok:
            self.e_date_return.delete(0, 'end')
            self.update_stock(); self.update_issue_cb()
            self.update_return_cb(); self.update_remove_cb()
            self.update_terms_cb()
            self.update_history_table()


    def cmd_remove(self):
        sel = self.cb_remove.get()
        if not sel:
            self.lbl_rem.config(text="Selecione aparelho.")
            return
        pid, _ = sel.split(' - ', 1)
        pid = int(pid)
        item = self.inv.find(pid)
        if item['status'] != 'Disponível':
            self.lbl_rem.config(text="Não pode remover produto emprestado.", foreground='red')
            return
        pwd = simpledialog.askstring("Autorização", "Digite a senha mãe:", show='*')
        if pwd != ADMIN_PASS:
            self.lbl_rem.config(text="Senha incorreta. Ação não autorizada.", foreground='red')
            return
        ok, msg = self.inv.remove(pid, ADMIN_USER)
        self.lbl_rem.config(text=msg, foreground='green' if ok else 'red')
        if ok:
            self.update_stock(); self.update_issue_cb()
            self.update_return_cb(); self.update_remove_cb()

    def cmd_search(self):
        term = self.e_search.get().strip()
        for i in self.tree_search.get_children():
            self.tree_search.delete(i)
        if not term:
            messagebox.showwarning("Busca", "Digite um termo para buscar.")
            return
        res = self.inv.search(term)
        if not res:
            messagebox.showinfo("Busca", "Nenhum resultado.")
            return
        for p in res:
            row = (
                p['id'], p['brand'], p['model'], p['identificador'], p['status'],
                p['assigned_to'] or "-", p.get('date_registered') or "-",
                p.get('date_issued') or "-", p.get('date_returned') or "-"
            )
            self.tree_search.insert('', 'end', values=row)
        self.autosize_treeview(self.tree_search)

    def cmd_generate_report(self):
        for i in self.tree_report.get_children():
            self.tree_report.delete(i)
        try:
            year = int(self.e_report_year.get())
            month = int(self.cb_report_month.get())
        except ValueError:
            messagebox.showerror("Erro", "Ano/Mês inválidos.")
            return

        flat_logs = self.inv.generate_monthly_report(year, month)
        for log in flat_logs:
            item = self.inv.find(log['item_id']) or {}
            row = (
                log['item_id'],
                item.get('tipo', '-'),
                item.get('brand', '-'),
                item.get('model', '-'),
                item.get('identificador', '-'),
                log.get('user', '-'),
                format_cpf(log.get('cpf', item.get('cpf', '-'))),
                log['operation'],
                format_date(log['date']),
                log.get('center_cost', '-'),
                log.get('cargo', '-'),
                log.get('revenda', '-')
            )
            self.tree_report.insert('', 'end', values=row)
        self.autosize_treeview(self.tree_report)


    def cmd_delete_report_entry(self):
        selected = self.tree_report.selection()
        if not selected:
            messagebox.showwarning("Aviso", "Selecione um lançamento do relatório.")
            return

        pwd = simpledialog.askstring("Autorização", "Digite a senha mãe:", show='*')
        if pwd != ADMIN_PASS:
            messagebox.showerror("Erro", "Senha incorreta.")
            return

        values = self.tree_report.item(selected[0])['values']
        # Mapeamento das colunas no relatório:
        # 0 ID | 1 Tipo | 2 Marca | 3 Modelo | 4 Identificador |
        # 5 Usuário | 6 CPF | 7 Operação | 8 Data | 9 Centro de Custo | 10 Cargo | 11 Revenda
        item_id = int(values[0])
        selected_user = values[5]
        selected_operation = values[7]
        selected_date_str = values[8]  # está em dd/mm/aaaa
        try:
            selected_date_iso = datetime.strptime(selected_date_str, "%d/%m/%Y").strftime("%Y-%m-%d")
        except:
            messagebox.showerror("Erro", "Data inválida na seleção.")
            return

        # Localiza o registro no histórico "completo"
        idx_encontrado = None
        for i, h in enumerate(self.inv.history):
            if h.get("id") != item_id:
                continue
            if selected_operation == "Empréstimo" and h.get("data_emprestimo") == selected_date_iso and h.get("usuario") == selected_user:
                idx_encontrado = i
                break
            if selected_operation == "Devolução" and h.get("data_devolucao") == selected_date_iso and h.get("usuario") == selected_user:
                idx_encontrado = i
                break

        if idx_encontrado is None:
            messagebox.showerror("Erro", "Lançamento não encontrado no histórico.")
            return

        h = self.inv.history[idx_encontrado]
        if selected_operation == "Empréstimo":
            # só pode estornar empréstimo se NÃO houver devolução nesse mesmo registro
            if h.get("data_devolucao"):
                messagebox.showerror("Erro", "Estorne primeiro a devolução deste empréstimo.")
                return
            # remove o registro inteiro
            del self.inv.history[idx_encontrado]
        else:
            # Estornar devolução => apenas limpar a data_devolucao
            h["data_devolucao"] = ""

        self.inv.save_history()
        self.inv.recalc_item_status(item_id)

        # Atualiza telas
        self.cmd_generate_report()
        self.update_stock()
        self.update_issue_cb()
        self.update_return_cb()
        self.update_remove_cb()
        self.update_terms_cb()
        self.update_history_table()

        messagebox.showinfo("Sucesso", "Lançamento estornado.")


    def cmd_generate_term(self):
        selected = self.tree_terms.selection()
        if not selected:
            self.lbl_terms.config(text="Selecione um empréstimo.", foreground='red')
            return
        values = self.tree_terms.item(selected[0])["values"]
        pid = int(values[0])
        user = values[3]  # coluna "Usuário"
        ok, filename = self.inv.generate_term(pid, user)
        if ok:
            self.lbl_terms.config(text=f"Termo gerado: {filename}", foreground='green')
            os.startfile(filename)
        else:
            self.lbl_terms.config(text=filename, foreground='red')



    # --- Funções de Gráficos ---
    def graph_issue_return(self):
        # lê ano/mês selecionados
        try:
            year = int(self.e_graph_year.get())
            month = int(self.cb_graph_month.get())
        except ValueError:
            messagebox.showerror("Erro", "Ano/Mês inválidos.")
            return

        issues, returns = self.inv.get_issue_return_counts()
        key = f"{year:04d}-{month:02d}"
        issue_val  = issues.get(key, 0)
        return_val = returns.get(key, 0)

        # gráfico de barras
        plt.figure(figsize=(6, 4))
        bars = plt.bar(["Empréstimos", "Devoluções"], [issue_val, return_val], color=['#4C72B0', '#55A868'])
        plt.title(f"Empréstimos vs Devoluções em {key}")
        plt.ylabel("Quantidade")
        for bar in bars:
            yval = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2, yval + 0.1, int(yval), ha='center', va='bottom')
        plt.tight_layout()
        plt.show()

    def graph_registration(self):
        # lê ano/mês selecionados
        try:
            year = int(self.e_graph_year.get())
            month = int(self.cb_graph_month.get())
        except ValueError:
            messagebox.showerror("Erro", "Ano/Mês inválidos.")
            return

        regs = self.inv.get_registration_counts()
        key = f"{year:04d}-{month:02d}"
        reg_val = regs.get(key, 0)

        plt.figure(figsize=(4, 4))
        bars = plt.bar([key], [reg_val], color='skyblue')
        plt.title(f"Cadastros em {key}")
        plt.ylabel("Total de Cadastros")
        for bar in bars:
            yval = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2, yval + 0.1, int(yval), ha='center', va='bottom')
        plt.tight_layout()
        plt.show()

if __name__ == "__main__":
    App().mainloop()