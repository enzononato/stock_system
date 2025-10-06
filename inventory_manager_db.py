import os
import shutil
from datetime import datetime
import pymysql
from docx import Document
import calendar

from database_mysql import get_connection
from config import TERMS_DIR, TERMO_MODELOS, REMOVAL_NOTES_DIR, SIGNED_TERMS_DIR, RETURN_TERMS_DIR, SIGNED_RETURN_TERMS_DIR, TERMO_DEVOLUCAO_MODELOS
from utils import format_cpf, format_date


class InventoryDBManager:
    def __init__(self):
        self._create_tables()


    def _create_tables(self):
        """Cria as tabelas no banco, caso ainda não existam"""
        conn = get_connection()
        cur = conn.cursor()
        
        # Tabela 'items'
        cur.execute("""
        CREATE TABLE IF NOT EXISTS items (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tipo VARCHAR(50), brand VARCHAR(100), model VARCHAR(100), identificador VARCHAR(100),
            nota_fiscal VARCHAR(50) UNIQUE,
            status ENUM('Disponível','Indisponível', 'Pendente', 'Pendente Devolução') DEFAULT 'Disponível',
            assigned_to VARCHAR(100), cpf VARCHAR(20), revenda VARCHAR(100),
            dominio VARCHAR(50), host VARCHAR(100), endereco_fisico VARCHAR(150),
            cpu VARCHAR(100), ram VARCHAR(50), storage VARCHAR(50),
            sistema VARCHAR(100), licenca VARCHAR(100), anydesk VARCHAR(50),
            setor VARCHAR(100), ip VARCHAR(50), mac VARCHAR(50),
            poe ENUM('Sim', 'Não'),
            quantidade_portas VARCHAR(10),
            date_registered DATETIME NOT NULL, date_issued DATETIME,
            is_active TINYINT(1) DEFAULT 1
        )
        """)
        
        # Tabela 'history'
        cur.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id INT AUTO_INCREMENT PRIMARY KEY,
            item_id INT,
            peripheral_id INT,
            operador VARCHAR(100),
            usuario VARCHAR(100),
            cpf VARCHAR(20),
            cargo VARCHAR(100),
            center_cost VARCHAR(100),
            revenda VARCHAR(100),
            data_operacao DATETIME,
            operation ENUM(
                'Cadastro','Empréstimo','Devolução', 'Edição', 'Exclusão', 'Estorno',
                'Confirmação Empréstimo', 'Confirmação Devolução',
                'Cadastro Periférico', 'Vínculo Periférico', 'Desvínculo Periférico', 'Substituição Periférico'
            ) DEFAULT 'Cadastro',
            is_reversed TINYINT(1) DEFAULT 0,
            details VARCHAR(255),
            tipo VARCHAR(50), brand VARCHAR(100), model VARCHAR(100), identificador VARCHAR(100),
            nota_fiscal VARCHAR(50), poe ENUM('Sim', 'Não'), quantidade_portas VARCHAR(10),
            nota_fiscal_anexo VARCHAR(255) NULL, termo_assinado_anexo VARCHAR(255) NULL,
            FOREIGN KEY(item_id) REFERENCES items(id) ON DELETE SET NULL
        )
        """)

        # --- TABELA DE PERIFÉRICOS
        cur.execute("""
        CREATE TABLE IF NOT EXISTS peripherals (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tipo VARCHAR(50) NOT NULL,
            brand VARCHAR(100),
            model VARCHAR(100),
            identificador VARCHAR(100) UNIQUE,
            status ENUM('Disponível', 'Em Uso', 'Com Defeito') DEFAULT 'Disponível',
            date_registered DATETIME NOT NULL,
            is_active TINYINT(1) DEFAULT 1
        )
        """)

        # --- TABELA DE VÍNCULO
        cur.execute("""
        CREATE TABLE IF NOT EXISTS equipment_peripherals (
            id INT AUTO_INCREMENT PRIMARY KEY,
            equipment_id INT NOT NULL,
            peripheral_id INT NOT NULL,
            FOREIGN KEY(equipment_id) REFERENCES items(id) ON DELETE CASCADE,
            FOREIGN KEY(peripheral_id) REFERENCES peripherals(id) ON DELETE CASCADE,
            UNIQUE(equipment_id, peripheral_id)
        )
        """)

        conn.commit()
        cur.close()
        conn.close()


    def add_item(self, item_data: dict, logged_user: str):
        """Adiciona novo item no estoque"""
        keys = ", ".join(item_data.keys())
        placeholders = ", ".join(["%s"] * len(item_data))
        values = list(item_data.values())
        sql = f"INSERT INTO items ({keys}) VALUES ({placeholders})"
        
        conn = get_connection()
        cur = conn.cursor()
        cur.execute(sql, values)
        item_id = cur.lastrowid

        cur.execute("""
            INSERT INTO history (item_id, operador, data_operacao, operation)
            VALUES (%s, %s, %s, 'Cadastro')
        """, (item_id, logged_user, datetime.now()))
        
        conn.commit()
        cur.close()
        conn.close()
        return item_id

    def update_item(self, item_id: int, item_data: dict, logged_user: str):
        """Atualiza os dados de um item e LOGA A AÇÃO"""
        sets = ", ".join([f"{k}=%s" for k in item_data.keys()])
        values = list(item_data.values())
        values.append(item_id)
        sql = f"UPDATE items SET {sets} WHERE id=%s"

        conn = get_connection()
        cur = conn.cursor()
        cur.execute(sql, values)

        cur.execute("""
            INSERT INTO history (item_id, operador, data_operacao, operation)
            VALUES (%s, %s, %s, 'Edição')
        """, (item_id, logged_user, datetime.now()))

        conn.commit()
        cur.close()
        conn.close()
        return True, f"Item {item_id} atualizado com sucesso."


    def remove(self, item_id: int, logged_user: str, attachment_path: str):
        """"Remove" (desativa) item do estoque, salvando seus dados no histórico."""
        item = self.find(item_id)
        if not item: return False, "ID não encontrado."
        if item["status"] != "Disponível": return False, "Não é possível remover produto emprestado."
        
        # --- Lógica para copiar o anexo ---
        try:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            original_filename = os.path.basename(attachment_path)
            # Cria um nome de arquivo seguro e único
            new_filename = f"remocao_{item_id}_{timestamp}_{original_filename}"
            destination_path = os.path.join(REMOVAL_NOTES_DIR, new_filename)
            
            # Copia o arquivo para a pasta de notas
            shutil.copy(attachment_path, destination_path)
        except Exception as e:
            return False, f"Erro ao processar o anexo: {e}"
        # ------------------------------------

        conn = get_connection()
        cur = conn.cursor()

        cur.execute("""
            INSERT INTO history (item_id, operador, data_operacao, operation, tipo, brand, model, identificador, nota_fiscal, nota_fiscal_anexo, poe, quantidade_portas)
            VALUES (%s, %s, %s, 'Exclusão', %s, %s, %s, %s, %s, %s, %s, %s)
        """, (item_id, logged_user, datetime.now(),
              item.get('tipo'), item.get('brand'), item.get('model'), item.get('identificador'), item.get('nota_fiscal'), destination_path,
              item.get('poe'), item.get('quantidade_portas')))

        cur.execute("UPDATE items SET is_active = 0 WHERE id=%s", (item_id,))
        
        conn.commit()
        cur.close()
        conn.close()
        return True, f"Aparelho {item_id} removido do estoque."


    def list_items(self):
        """Lista os itens e conta quantos periféricos estão vinculados a cada um."""
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        cur.execute("""
            SELECT i.*, COUNT(ep.peripheral_id) as peripheral_count
            FROM items i
            LEFT JOIN equipment_peripherals ep ON i.id = ep.equipment_id
            WHERE i.is_active = 1
            GROUP BY i.id
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows


    def find(self, item_id: int):
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        cur.execute("SELECT * FROM items WHERE id=%s AND is_active = 1", (item_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return row
    

    # --- NOVAS FUNÇÕES DE GERENCIAMENTO DE PERIFÉRICOS ---

# Em inventory_manager_db.py

    def add_peripheral(self, data: dict, logged_user: str):
        """Adiciona um novo periférico."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            
            data['date_registered'] = datetime.now() 
            keys = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            sql = f"INSERT INTO peripherals ({keys}) VALUES ({placeholders})"
            cur.execute(sql, list(data.values()))
            p_id = cur.lastrowid

            cur.execute("""
                INSERT INTO history (peripheral_id, operador, data_operacao, operation, tipo, brand, model, identificador)
                VALUES (%s, %s, %s, 'Cadastro Periférico', %s, %s, %s, %s)
            """, (p_id, logged_user, datetime.now(), data.get('tipo'), data.get('brand'), data.get('model'), data.get('identificador')))
            conn.commit()
            return True, f"Periférico '{data.get('tipo')}' cadastrado com ID {p_id}."
        except pymysql.MySQLError as e:
            conn.rollback()
            if e.args[0] == 1062:
                return False, "Erro: Já existe um periférico com este Identificador (Nº de Série)."
            return False, f"Erro de banco de dados: {e}"
        finally:
            cur.close()
            conn.close()

    def list_peripherals(self, status_filter="", type_filter=""):
        """Lista os periféricos ATIVOS com filtros opcionais."""
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)

        query = "SELECT * FROM peripherals WHERE is_active = 1"
        params = []
        if status_filter:
            query += " AND status = %s"
            params.append(status_filter)
        if type_filter:
            query += " AND tipo = %s"
            params.append(type_filter)
        
        cur.execute(query, params)
        return cur.fetchall()

    def list_peripherals_for_equipment(self, equipment_id: int):
        """Lista os periféricos vinculados a um equipamento específico."""
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT p.*, ep.id as link_id
            FROM peripherals p
            JOIN equipment_peripherals ep ON p.id = ep.peripheral_id
            WHERE ep.equipment_id = %s
        """, (equipment_id,))
        return cur.fetchall()

    def link_peripheral_to_equipment(self, equipment_id: int, peripheral_id: int, logged_user: str):
        """Cria um vínculo e atualiza o status do periférico."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            
            cur.execute("INSERT INTO equipment_peripherals (equipment_id, peripheral_id) VALUES (%s, %s)", (equipment_id, peripheral_id))
            cur.execute("UPDATE peripherals SET status = 'Em Uso' WHERE id = %s", (peripheral_id,))
            cur.execute("""
                INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation)
                VALUES (%s, %s, %s, %s, 'Vínculo Periférico')
            """, (equipment_id, peripheral_id, logged_user, datetime.now()))
            conn.commit()
            return True, "Periférico vinculado com sucesso."
        except pymysql.MySQLError as e:
            conn.rollback()
            return False, f"Erro ao vincular: {e}"
        finally:
            cur.close()
            conn.close()

    def unlink_peripheral_from_equipment(self, link_id: int, logged_user: str):
        """Remove um vínculo e atualiza o status do periférico."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            
            # Pega os IDs antes de deletar para poder registrar no histórico
            cur.execute("SELECT equipment_id, peripheral_id FROM equipment_peripherals WHERE id = %s", (link_id,))
            ids = cur.fetchone()
            if not ids:
                return False, "Vínculo não encontrado."

            cur.execute("DELETE FROM equipment_peripherals WHERE id = %s", (link_id,))
            cur.execute("UPDATE peripherals SET status = 'Disponível' WHERE id = %s", (ids['peripheral_id'],))
            cur.execute("""
                INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation)
                VALUES (%s, %s, %s, %s, %s)
            """, (ids['equipment_id'], ids['peripheral_id'], logged_user, datetime.now(), 'Desvínculo Periférico'))
            conn.commit()
            return True, "Periférico desvinculado com sucesso."
        except pymysql.MySQLError as e:
            conn.rollback()
            return False, f"Erro ao desvincular: {e}"
        finally:
            cur.close()
            conn.close()


    def replace_peripheral(self, equipment_id: int, old_peripheral_id: int, new_peripheral_id: int, reason: str, logged_user: str):
        """Substitui um periférico por outro."""
        conn = get_connection()
        cur = conn.cursor()
        try:
            # 1. Marca o periférico antigo como "Com Defeito" E INATIVO
            cur.execute("UPDATE peripherals SET status = 'Com Defeito', is_active = 0 WHERE id = %s", (old_peripheral_id,))
            
            # 2. Remove o vínculo antigo
            cur.execute("DELETE FROM equipment_peripherals WHERE equipment_id = %s AND peripheral_id = %s", (equipment_id, old_peripheral_id))
            # 3. Cria o novo vínculo
            cur.execute("INSERT INTO equipment_peripherals (equipment_id, peripheral_id) VALUES (%s, %s)", (equipment_id, new_peripheral_id))
            # 4. Marca o novo periférico como "Em Uso"
            cur.execute("UPDATE peripherals SET status = 'Em Uso' WHERE id = %s", (new_peripheral_id,))
            # 5. Registra a substituição no histórico
            details_log = f"Substituído periférico ID {old_peripheral_id} por ID {new_peripheral_id}. Motivo: {reason}"
            cur.execute("""
                INSERT INTO history (item_id, operador, data_operacao, operation, details)
                VALUES (%s, %s, %s, 'Substituição Periférico', %s)
            """, (equipment_id, logged_user, datetime.now(), details_log))
            conn.commit()
            return True, "Substituição realizada com sucesso."
        except pymysql.MySQLError as e:
            conn.rollback()
            return False, f"Erro ao substituir periférico: {e}"
        finally:
            cur.close()
            conn.close()


    def issue(self, pid, user, cpf, center_cost, cargo, revenda, date_issue, logged_user: str):
        item = self.find(pid)
        if not item:
            return False, "Item não encontrado."

        if item['status'] != 'Disponível':
            return False, f"Este item não está disponível para empréstimo."
        
        try:
            dt_issue = datetime.strptime(date_issue, "%d/%m/%Y")
            # Mantém a hora atual
            dt_issue = dt_issue.replace(
            hour=datetime.now().hour,
            minute=datetime.now().minute,
            second=datetime.now().second
        )
        except ValueError:
            return False, "Data de empréstimo inválida (use dd/mm/aaaa)."
        
        if dt_issue.date() > datetime.now().date():
            return False, "A data de empréstimo não pode ser no futuro."

        if item.get('date_registered') and dt_issue.date() < item['date_registered'].date():
            data_formatada = format_date(str(item['date_registered']))
            return False, f"Data de empréstimo não pode ser anterior ao cadastro ({data_formatada})."

        conn = get_connection()
        cur = conn.cursor()
        #  O status agora é 'Pendente'
        cur.execute("""
            UPDATE items 
            SET status='Pendente', assigned_to=%s, cpf=%s, date_issued=%s, revenda=%s 
            WHERE id=%s
        """, (user, cpf, dt_issue, revenda, pid))

        cur.execute("""
            INSERT INTO history (item_id, operador, usuario, cpf, cargo, center_cost, revenda, data_operacao, operation)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'Empréstimo')
        """, (pid, logged_user, user, cpf, cargo, center_cost, revenda, dt_issue, 'Empréstimo'))

        conn.commit()
        cur.close()
        conn.close()
        return True, f"Empréstimo do item {pid} para {user} iniciado. Status: Pendente."
    

    def confirm_loan(self, item_id: int, logged_user: str, signed_term_path: str):
        """
        Confirma o empréstimo, anexa o termo, muda o status do item E DOS PERIFÉRICOS VINCULADOS.
        """
        item = self.find(item_id)
        if not item or item['status'] != 'Pendente':
            return False, "Apenas itens com status 'Pendente' podem ser confirmados."

        try:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            original_filename = os.path.basename(signed_term_path)
            new_filename = f"termo_{item_id}_{timestamp}_{original_filename}"
            destination_path = os.path.join(SIGNED_TERMS_DIR, new_filename)
            shutil.copy(signed_term_path, destination_path)
        except Exception as e:
            return False, f"Erro ao processar o termo assinado: {e}"

        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        try:
            

            #Pega os detalhes do empréstimo original no histórico
            cur.execute("""
                SELECT id, usuario, cpf, cargo, center_cost, revenda FROM history
                WHERE item_id = %s AND operation = 'Empréstimo' AND is_reversed = 0
                ORDER BY data_operacao DESC, id DESC LIMIT 1
            """, (item_id,))
            last_loan_details = cur.fetchone()
            if not last_loan_details:
                raise Exception("Registro de empréstimo original não encontrado no histórico.")

            # ATUALIZA O STATUS DO ITEM PRINCIPAL
            cur.execute("UPDATE items SET status='Indisponível' WHERE id=%s", (item_id,))

            # ATUALIZA O STATUS DOS PERIFÉRICOS VINCULADOS
            linked_peripherals = self.list_peripherals_for_equipment(item_id)
            for p in linked_peripherals:
                cur.execute("UPDATE peripherals SET status='Em Uso' WHERE id=%s", (p['id'],))

            # Cria o novo registro de "Confirmação" no histórico
            cur.execute("""
                INSERT INTO history (item_id, operador, usuario, cpf, cargo, center_cost, revenda, data_operacao, operation, termo_assinado_anexo)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'Confirmação Empréstimo', %s)
            """, (
                item_id, logged_user,
                last_loan_details.get('usuario'), last_loan_details.get('cpf'),
                last_loan_details.get('cargo'), last_loan_details.get('center_cost'),
                last_loan_details.get('revenda'),
                datetime.now(),
                destination_path
            ))
            
            # cconfirma todas as operações no banco
            conn.commit()
            return True, f"Empréstimo do item {item_id} e de {len(linked_peripherals)} periférico(s) confirmado."
        except Exception as e:
            # Desfaz tudo se der algum erro no meio do caminho
            conn.rollback()
            return False, f"Erro ao confirmar empréstimo: {e}"
        finally:
            cur.close()
            conn.close()
    
    
# PROCESSO DE DEVOLUÇÃO, 2 PRÓXIMAS FUNÇÕES
    
    def generate_and_initiate_return(self, item_id: int, logged_user: str):
        """
        Gera o termo de devolução e, se bem-sucedido, inicia o processo
        mudando o status do item para 'Pendente Devolução'.
        """
        item = self.find(item_id)
        if not item or item['status'] != 'Indisponível':
            return False, "Apenas itens com status 'Indisponível' podem ter um termo de devolução gerado."

        # Pega o nome do usuário que está com o item
        user = item.get("assigned_to")
        if not user:
            return False, "Não foi possível encontrar o usuário associado a este empréstimo."

        # --- Parte 1: Gerar o Termo (lógica de generate_return_term) ---
        revenda = item.get("revenda")
        modelo_path = TERMO_DEVOLUCAO_MODELOS.get(revenda)
        if not modelo_path or not os.path.exists(modelo_path):
            return False, f"Modelo de termo de devolução não encontrado para {revenda}."

        safe_user_name = user.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saida_path = os.path.join(RETURN_TERMS_DIR, f"termo_devolucao_{item_id}_{safe_user_name}_{revenda}_{timestamp}.docx")
        doc = Document(modelo_path)
        
        substituicoes = {
            "{{nome}}": user, "{{data_hoje}}": datetime.now().strftime("%d/%m/%Y"),
            "{{cpf}}": format_cpf(item.get("cpf", "")),
            "{{data_emprestimo}}": format_date(item.get("date_issued", "")), "{{marca}}": f" {item.get('brand', '')}" if item.get("brand") else "",
            "{{modelo}}": f" {item.get('model', '')}" if item.get("model") else "", "{{tipo}}": item.get("tipo", ""),
            "{{identificador}}": f" {item.get('identificador', '')}" if item.get("identificador") else "",
            "{{nota_fiscal}}": f" {item.get('nota_fiscal', '')}" if item.get("nota_fiscal") else ""
        }
        
        try:
            for p in doc.paragraphs:
                for chave, valor in substituicoes.items():
                    if chave in p.text: p.text = p.text.replace(chave, str(valor))
            
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for p in celula.paragraphs:
                            for chave, valor in substituicoes.items():
                                if chave in p.text: p.text = p.text.replace(chave, str(valor))
            doc.save(saida_path)
        except Exception as e:
            return False, f"Erro ao salvar o documento do termo: {e}"

        # --- Parte 2: Iniciar a Devolução (lógica de initiate_return) ---
        conn = get_connection()
        cur = conn.cursor()
        
        cur.execute("UPDATE items SET status='Pendente Devolução' WHERE id=%s", (item_id,))
        
        cur.execute("""
            INSERT INTO history (item_id, operador, data_operacao, operation)
            VALUES (%s, %s, %s, 'Devolução')
        """, (item_id, logged_user, datetime.now()))

        conn.commit()
        cur.close()
        conn.close()
        
        # Se tudo deu certo, retorna True e o caminho do arquivo gerado
        return True, saida_path



    def confirm_return(self, item_id: int, logged_user: str, signed_return_term_path: str):
        """
        Finaliza a devolução, anexa o termo, muda o status do item E DOS PERIFÉRICOS VINCULADOS.
        """
        item = self.find(item_id)
        if not item or item['status'] != 'Pendente Devolução':
            return False, "Apenas itens com status 'Pendente Devolução' podem ser confirmados."

        try:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            original_filename = os.path.basename(signed_return_term_path)
            new_filename = f"termo_devolucao_{item_id}_{timestamp}_{original_filename}"
            destination_path = os.path.join(SIGNED_RETURN_TERMS_DIR, new_filename)
            shutil.copy(signed_return_term_path, destination_path)
        except Exception as e:
            return False, f"Erro ao processar o termo de devolução assinado: {e}"

        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        try:
            

            # ATUALIZA O STATUS DO ITEM PRINCIPAL
            cur.execute("""
                UPDATE items SET status='Disponível', assigned_to=NULL, cpf=NULL, date_issued=NULL 
                WHERE id=%s
            """, (item_id,))

            #ATUALIZA O STATUS DOS PERIFÉRICOS VINCULADOS
            linked_peripherals = self.list_peripherals_for_equipment(item_id)
            for p in linked_peripherals:
                cur.execute("UPDATE peripherals SET status='Disponível' WHERE id=%s", (p['id'],))

            # Cria o novo registro de "Confirmação Devolução" no histórico
            cur.execute("""
                INSERT INTO history (item_id, operador, data_operacao, operation, termo_assinado_anexo)
                VALUES (%s, %s, %s, 'Confirmação Devolução', %s)
            """, (
                item_id,
                logged_user,
                datetime.now(), #Salva data e hora completas
                destination_path
            ))

            conn.commit()
            return True, f"Devolução do item {item_id} e de {len(linked_peripherals)} periférico(s) confirmada."
        except Exception as e:
            # Desfaz tudo se der erro
            conn.rollback()
            return False, f"Erro ao confirmar devolução: {e}"
        finally:
            cur.close()
            conn.close()


    def list_history(self):
        """Lista histórico de TODAS as operações, tratando itens excluídos."""
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        # ALTERADO: Adicionado filtro para não mostrar operações estornadas
        cur.execute("""
            SELECT
                h.id, h.item_id, h.operador, h.peripheral_id, h.details,
                    
                -- Usa COALESCE para pegar dados do item ou do histórico, se o item foi excluído
                COALESCE(i.tipo, h.tipo) as tipo,
                COALESCE(i.brand, h.brand) as marca,
                COALESCE(i.model, h.model) as modelo,
                COALESCE(i.identificador, h.identificador) as identificador,
                COALESCE(i.nota_fiscal, h.nota_fiscal) as nota_fiscal,
                h.usuario, h.cpf, h.cargo, h.center_cost, h.revenda,
                h.data_operacao, h.operation
            FROM history h
            LEFT JOIN items i ON i.id = h.item_id
            WHERE h.is_reversed = 0
            ORDER BY h.data_operacao DESC, h.id DESC
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows



    def generate_monthly_report(self, ano, mes):
        """Relatório consolidado de um mês, que agora diferencia empréstimos pendentes, confirmados e devolvidos."""
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        
        # A nova query usa subconsultas para encontrar as datas de confirmação e devolução
        # associadas a cada registro de empréstimo.
        sql = """
            WITH EmprestimosDoMes AS (
                SELECT
                    h.id AS history_id,
                    h.item_id, h.operador, h.usuario, h.cpf, h.cargo, h.center_cost, h.revenda,
                    h.data_operacao AS data_emprestimo,
                    
                    (SELECT MIN(hc.data_operacao) FROM history hc 
                     WHERE hc.item_id = h.item_id AND hc.operation = 'Confirmação Empréstimo' 
                     AND hc.id > h.id AND hc.is_reversed = 0) AS data_confirmacao,
                     
                    (SELECT MIN(hd.data_operacao) FROM history hd 
                     WHERE hd.item_id = h.item_id AND hd.operation = 'Devolução' 
                     AND hd.id > h.id AND hd.is_reversed = 0) AS data_devolucao
                FROM
                    history h
                WHERE
                    h.operation = 'Empréstimo'
                    AND h.is_reversed = 0
                    AND YEAR(h.data_operacao) = %s AND MONTH(h.data_operacao) = %s
            )
            SELECT
                em.*,
                'Empréstimo' AS operation_type,
                COALESCE(i.tipo, h_item.tipo) AS tipo,
                COALESCE(i.brand, h_item.brand) AS brand,
                COALESCE(i.model, h_item.model) AS model,
                COALESCE(i.identificador, h_item.identificador) AS identificador,
                COALESCE(i.nota_fiscal, h_item.nota_fiscal) AS nota_fiscal
            FROM EmprestimosDoMes em
            LEFT JOIN items i ON i.id = em.item_id
            LEFT JOIN history h_item ON h_item.id = em.history_id

            UNION ALL

            SELECT
                cad.id as history_id, cad.item_id, cad.operador, NULL, NULL, NULL, NULL, i.revenda,
                cad.data_operacao AS data_emprestimo,
                NULL AS data_confirmacao,
                NULL AS data_devolucao,
                'Cadastro' AS operation_type,
                COALESCE(i.tipo, cad.tipo) AS tipo,
                COALESCE(i.brand, cad.brand) AS brand,
                COALESCE(i.model, cad.model) AS model,
                COALESCE(i.identificador, cad.identificador) AS identificador,
                COALESCE(i.nota_fiscal, cad.nota_fiscal) AS nota_fiscal
            FROM history cad
            LEFT JOIN items i ON i.id = cad.item_id
            WHERE cad.operation = 'Cadastro'
            AND cad.is_reversed = 0
            AND YEAR(cad.data_operacao) = %s AND MONTH(cad.data_operacao) = %s
            
            ORDER BY data_emprestimo, item_id;
        """
        cur.execute(sql, (ano, mes, ano, mes))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows

    # ALTERADO: Lógica de estorno completamente refeita
    def reverse_history_entry(self, history_id: int, logged_user: str):
        """
        Estorna um lançamento: reverte o estado do item, marca a operação original
        como 'estornada' e cria um novo registro de 'Estorno' para auditoria.
        """
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        
        try:

            cur.execute("SELECT * FROM history WHERE id = %s", (history_id,))
            entry_to_reverse = cur.fetchone()

            if not entry_to_reverse:
                return False, "Lançamento do histórico não encontrado."
            
            if entry_to_reverse['is_reversed']:
                return False, "Esta operação já foi estornada anteriormente."

            op = entry_to_reverse['operation']
            item_id = entry_to_reverse['item_id']

            # --- LÓGICA DE REVERSÃO ATUALIZADA ---
            if op == 'Confirmação Empréstimo':
                cur.execute("UPDATE items SET status='Pendente' WHERE id = %s", (item_id,))
            
            elif op == 'Empréstimo':
                cur.execute("UPDATE items SET status='Disponível', assigned_to=NULL, cpf=NULL, date_issued=NULL WHERE id = %s", (item_id,))
            
            elif op == 'Confirmação Devolução':
                # Reverte o item para 'Pendente Devolução', mas mantém os dados do usuário
                # pois ele ainda está "associado" ao item até a devolução ser estornada por completo.
                cur.execute("UPDATE items SET status='Pendente Devolução' WHERE id = %s", (item_id,))
            
            elif op == 'Devolução':
                # Encontra os dados do último empréstimo válido para restaurar o estado
                cur.execute("""
                    SELECT usuario, cpf, data_operacao FROM history
                    WHERE item_id = %s AND operation IN ('Empréstimo', 'Confirmação Empréstimo') AND id < %s AND is_reversed = 0
                    ORDER BY data_operacao DESC, id DESC LIMIT 1
                """, (item_id, history_id))
                last_loan = cur.fetchone()

                if not last_loan:
                    conn.rollback()
                    return False, "Não foi possível encontrar o empréstimo original para reverter a devolução."

                cur.execute("UPDATE items SET status='Indisponível', assigned_to=%s, cpf=%s, date_issued=%s WHERE id = %s", 
                            (last_loan['usuario'], last_loan['cpf'], last_loan['data_operacao'], item_id))
            
            elif op == 'Cadastro':
                cur.execute("UPDATE items SET is_active = 0 WHERE id = %s", (item_id,))
            
            else:
                conn.rollback()
                return False, f"Não é possível estornar uma operação do tipo '{op}'."

            # Marca a operação original como estornada
            cur.execute("UPDATE history SET is_reversed = 1 WHERE id = %s", (history_id,))
            
            # Cria o novo registro de 'Estorno' para auditoria
            cur.execute("""
                INSERT INTO history (item_id, operador, data_operacao, operation, usuario, cpf, cargo, center_cost, revenda)
                VALUES (%s, %s, %s, 'Estorno', %s, %s, %s, %s, %s)
            """, (
                item_id, logged_user, datetime.now(),
                entry_to_reverse.get('usuario'), entry_to_reverse.get('cpf'),
                entry_to_reverse.get('cargo'), entry_to_reverse.get('center_cost'),
                entry_to_reverse.get('revenda')
            ))
            
            conn.commit()
            return True, f"Operação '{op}' do item {item_id} estornada com sucesso."

        except pymysql.MySQLError as e:
            conn.rollback()
            return False, f"Erro ao estornar: {e}"
        finally:
            cur.close()
            conn.close()


    def generate_term(self, item_id, user):
        """Gera o termo de responsabilidade, incluindo a lista de periféricos vinculados."""
        item = self.find(item_id)
        if not item: 
            return False, "Equipamento não encontrado."
        

        # Verifica por "Pendente", que é o estado correto para gerar um termo.
        assigned_user = str(item.get("assigned_to", "")).strip()
        user_param = str(user).strip()
        if item["status"] != "Pendente" or assigned_user != user_param:
            return False, "Este equipamento não está pendente de empréstimo para este usuário."
        # ----------------------------------------

        # --- LÓGICA PARA BUSCAR E FORMATAR PERIFÉRICOS ---
        linked_peripherals = self.list_peripherals_for_equipment(item_id)
        peripherals_text = "Nenhum periférico adicional vinculado."
        if linked_peripherals:
            peripherals_list = [
                f"- {p['tipo']}: {p.get('brand','')} {p.get('model','')} (S/N: {p.get('identificador', 'N/A')})"
                for p in linked_peripherals
            ]
            peripherals_text = "\n".join(peripherals_list)
        # --- FIM DA LÓGICA ---

        revenda = item.get("revenda")
        modelo_path = TERMO_MODELOS.get(revenda)
        if not modelo_path or not os.path.exists(modelo_path): 
            return False, f"Modelo de termo não encontrado para {revenda}."
            
        safe_user_name = user_param.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saida_path = os.path.join(TERMS_DIR, f"termo_{item_id}_{safe_user_name}_{revenda}_{timestamp}.docx")
        doc = Document(modelo_path)
        
        substituicoes = {
            "{{nome}}": user, "{{data_hoje}}": datetime.now().strftime("%d/%m/%Y"),
            "{{cpf}}": format_cpf(item.get("cpf", "")),
            "{{data_emprestimo}}": format_date(item.get("date_issued", "")),
            "{{tipo}}": item.get("tipo", ""),
            "{{marca}}": f" {item.get('brand', '')}" if item.get("brand") else "",
            "{{modelo}}": f" {item.get('model', '')}" if item.get("model") else "",
            "{{identificador}}": f" {item.get('identificador', '')}" if item.get("identificador") else "",
            "{{nota_fiscal}}": f" {item.get('nota_fiscal', '')}" if item.get("nota_fiscal") else "",
            "{{perifericos}}": peripherals_text
        }
        
        for key, value in item.items():
            placeholder = f"{{{{{key}}}}}"
            str_value = str(value) if value not in [None, ""] else ""
            if str_value: str_value = " " + str_value
            substituicoes[placeholder] = str_value
        
        for p in doc.paragraphs:
            for chave, valor in substituicoes.items():
                if chave in p.text: p.text = p.text.replace(chave, str(valor))
        
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        for chave, valor in substituicoes.items():
                            if chave in p.text: p.text = p.text.replace(chave, str(valor))
                            
        doc.save(saida_path)
        return True, saida_path

    # --- NOVAS FUNÇÕES PARA GRÁFICOS ---
    def get_issue_return_counts(self, year, month):
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        num_days = calendar.monthrange(year, month)[1]
        days = list(range(1, num_days + 1))
        issues = {day: 0 for day in days}
        returns = {day: 0 for day in days}

        # ALTERADO: Adicionado filtro para ignorar operações estornadas
        sql = """
            SELECT DAY(data_operacao) as dia, operation, COUNT(id) as total
            FROM history
            WHERE
                YEAR(data_operacao) = %s AND
                MONTH(data_operacao) = %s AND
                operation IN ('Empréstimo', 'Devolução') AND
                is_reversed = 0
            GROUP BY dia, operation
        """
        cur.execute(sql, (year, month))
        
        for row in cur.fetchall():
            if row['operation'] == 'Empréstimo':
                issues[row['dia']] = row['total']
            elif row['operation'] == 'Devolução':
                returns[row['dia']] = row['total']
        cur.close()
        conn.close()
        return list(issues.keys()), list(issues.values()), list(returns.values())

    def get_registration_counts(self, year, month):
        conn = get_connection()
        cur = conn.cursor(pymysql.cursors.DictCursor)
        num_days = calendar.monthrange(year, month)[1]
        days = list(range(1, num_days + 1))
        registrations = {day: 0 for day in days}

        # ALTERADO: Adicionado filtro para ignorar operações estornadas
        sql = """
            SELECT DAY(data_operacao) as dia, COUNT(id) as total
            FROM history
            WHERE
                YEAR(data_operacao) = %s AND
                MONTH(data_operacao) = %s AND
                operation = 'Cadastro' AND
                is_reversed = 0
            GROUP BY dia
        """
        cur.execute(sql, (year, month))
        
        for row in cur.fetchall():
            registrations[row['dia']] = row['total']
        cur.close()
        conn.close()
        return list(registrations.keys()), list(registrations.values())