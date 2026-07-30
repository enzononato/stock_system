"""
Testes de geração do termo de responsabilidade (Módulo 10).

Este arquivo é posse exclusiva do Módulo 10 e não usa nenhuma fixture nova em
conftest.py (que não é posse deste módulo) — tudo que os testes precisam está
definido aqui.

Independência de app.db.unidade_db
-----------------------------------
`app.db.unidade_db.UnidadeDBManager` é posse de outro módulo, desenvolvido em
paralelo. `InventoryDBManager.generate_loan_term_bytes` consome apenas o
contrato documentado:

    UnidadeDBManager.get_unidade_por_nome(nome: str) -> dict | None
    # {id, nome, razao_social, cnpj, endereco, cep, cidade, uf, is_active}

Para não depender de qual implementação real está mesclada neste worktree no
momento em que os testes rodam (nem do schema real de bancos de unidades), os
testes abaixo substituem (monkeypatch) a referência `UnidadeDBManager` vista
por `app.db.inventory_manager_db` por um dublê que só respeita o contrato
acima — nenhum teste aqui toca em banco de dados de verdade.

Independência de banco de dados (items/peripherals)
----------------------------------------------------
Pelo mesmo motivo, os testes de geração de documento também não usam
`InventoryDBManager` real (que abriria conexão MySQL em `__init__`/`find`):
usam uma subclasse local que sobrescreve `find`/`list_peripherals_for_equipment`
com dados fixos em memória. Isso mantém a suíte deste módulo executável mesmo
sem Docker/MySQL de teste no ar — o que importa aqui é a geração do .docx, não
a camada de persistência (já coberta por outros módulos de teste).
"""
import io
import re
from datetime import datetime

from docx import Document as DocxDocument

from app.db import inventory_manager_db as inv_mod

PLACEHOLDER_RE = re.compile(r"\{\{\s*[a-zA-Z_]+\s*\}\}")


# ────────────────────────────────────────────────────────────────────────────
# Dublês (sem banco de dados real)
# ────────────────────────────────────────────────────────────────────────────

def _instalar_unidade_fake(monkeypatch, resultado):
    """Substitui UnidadeDBManager visto por app.db.inventory_manager_db por um
    dublê que devolve exatamente `resultado` (dict ou None) de
    get_unidade_por_nome, respeitando só o contrato documentado."""

    class _UnidadeFake:
        def __init__(self):
            pass

        def get_unidade_por_nome(self, nome):
            return resultado

    monkeypatch.setattr(inv_mod, "UnidadeDBManager", _UnidadeFake)


class _FakeManager(inv_mod.InventoryDBManager):
    """InventoryDBManager sem __init__/find/list_peripherals_for_equipment reais
    (que exigiriam conexão MySQL) — mantém generate_loan_term_bytes e
    generate_return_term_bytes (métodos reais, não sobrescritos) exercitados
    sobre dados fixos em memória."""

    def __init__(self, item, peripherals=None):
        self._item = item
        self._peripherals = peripherals or []

    def find(self, item_id):
        return self._item

    def list_peripherals_for_equipment(self, item_id):
        return self._peripherals


def _item_padrao(**overrides):
    dados = {
        "id": 1,
        "tipo": "Notebook",
        "brand": "Dell",
        "model": "Latitude 5420",
        "identificador": "SN-12345",
        "status": "Pendente",
        "assigned_to": "Fulano de Tal",
        "cpf": "11122233344",
        "revenda": "Revalle Juazeiro",
        "cpu": "Intel i5-1135G7",
        "ram": "8GB",
        "storage": "256GB SSD",
        "sistema": "Windows 11 Pro",
        "licenca": "OEM",
        "host": "NB-FULANO",
        "ip": "10.0.0.5",
        "mac": "AA:BB:CC:DD:EE:FF",
        "dominio": "revalle.local",
        "anydesk": "123 456 789",
        "potencia_nominal": None,
        "autonomia_estimada": None,
        "ip_snmp": None,
        "poe": None,
        "quantidade_portas": None,
        "codigo_patrimonial": "PAT-001",
        "local_instalacao": "Sala de TI",
        "date_issued": datetime(2026, 1, 10),
    }
    dados.update(overrides)
    return dados


def _unidade_padrao(**overrides):
    dados = {
        "id": 1,
        "nome": "Revalle Juazeiro",
        "razao_social": "Revenda Valle da Integração LTDA",
        "cnpj": "12.345.678/0001-90",
        "endereco": "Rua Principal, 100, Centro",
        "cep": "48900-000",
        "cidade": "Juazeiro",
        "uf": "BA",
        "is_active": 1,
    }
    dados.update(overrides)
    return dados


# ────────────────────────────────────────────────────────────────────────────
# Utilitários de inspeção do .docx gerado
# ────────────────────────────────────────────────────────────────────────────

def _paragrafos_incluindo_cabecalho_rodape(doc):
    """Réplica independente (não reaproveita o código de produção) da varredura
    de parágrafos + tabelas + cabeçalhos/rodapés, para que um eventual esquecimento
    de alguma seção em `_substituir_placeholders_documento` seja pego por estes
    testes em vez de silenciosamente reaproveitar o mesmo blind spot."""

    def _de(container):
        for p in container.paragraphs:
            yield p
        for tabela in container.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    yield from _de(celula)

    yield from _de(doc)
    for secao in doc.sections:
        for container in (
            secao.header, secao.footer,
            secao.first_page_header, secao.first_page_footer,
            secao.even_page_header, secao.even_page_footer,
        ):
            yield from _de(container)


def _texto_completo(doc):
    return "\n".join(p.text for p in _paragrafos_incluindo_cabecalho_rodape(doc))


def _gerar_termo_emprestimo(monkeypatch, item, unidade, peripherals=None):
    _instalar_unidade_fake(monkeypatch, unidade)
    manager = _FakeManager(item, peripherals)
    return manager.generate_loan_term_bytes(item["id"])


# ────────────────────────────────────────────────────────────────────────────
# Geração do termo de empréstimo (template único + dados da unidade)
# ────────────────────────────────────────────────────────────────────────────

class TestGeracaoDoTermoDeEmprestimo:
    def test_termo_gerado_e_um_docx_valido(self, monkeypatch):
        ok, resultado, filename = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(), _unidade_padrao()
        )
        assert ok, resultado
        assert filename.endswith(".docx")
        # Não deve lançar: um DOCX corrompido faria Document() explodir aqui.
        doc = DocxDocument(io.BytesIO(resultado))
        assert doc.paragraphs

    def test_nenhum_placeholder_sobra_no_documento_final(self, monkeypatch):
        ok, resultado, _ = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(), _unidade_padrao()
        )
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))
        texto = _texto_completo(doc)
        sobras = PLACEHOLDER_RE.findall(texto)
        assert not sobras, f"Placeholders não substituídos no termo final: {sobras}"

    def test_dados_da_unidade_aparecem_no_texto(self, monkeypatch):
        unidade = _unidade_padrao(
            razao_social="Distribuidora Exemplo LTDA",
            cnpj="99.888.777/0001-66",
            cidade="Paulo Afonso",
            uf="BA",
        )
        ok, resultado, _ = _gerar_termo_emprestimo(monkeypatch, _item_padrao(), unidade)
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))
        texto = _texto_completo(doc)
        assert "Distribuidora Exemplo LTDA" in texto
        assert "99.888.777/0001-66" in texto
        assert "Paulo Afonso/BA" in texto

    def test_negrito_do_rotulo_tipo_de_equipamento_e_preservado(self, monkeypatch):
        """Regressão do bug corrigido pela T1: `paragraph.text = ...` colapsava
        todos os runs num único run sem formatação, e o rótulo em negrito
        "Tipo de Equipamento:" saía sem negrito no termo final."""
        ok, resultado, _ = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(), _unidade_padrao()
        )
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))

        rotulo_encontrado = False
        for p in doc.paragraphs:
            if "Tipo de Equipamento:" in p.text:
                for run in p.runs:
                    if "Tipo de Equipamento:" in run.text:
                        assert run.bold is True, (
                            "O run do rótulo 'Tipo de Equipamento:' perdeu o negrito "
                            "após a substituição de placeholders."
                        )
                        rotulo_encontrado = True
        assert rotulo_encontrado, "Parágrafo com 'Tipo de Equipamento:' não encontrado no termo gerado."

    def test_sem_periferico_e_sem_campo_tecnico_nao_deixa_placeholder_cru_nem_linha_vazia(
        self, monkeypatch
    ):
        item = _item_padrao(
            cpu=None, ram=None, storage=None, sistema=None, licenca=None,
            host=None, ip=None, mac=None, dominio=None, anydesk=None,
            potencia_nominal=None, autonomia_estimada=None, ip_snmp=None,
            poe=None, quantidade_portas=None, codigo_patrimonial=None,
            local_instalacao=None,
        )
        ok, resultado, _ = _gerar_termo_emprestimo(
            monkeypatch, item, _unidade_padrao(), peripherals=[]
        )
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))
        texto = _texto_completo(doc)

        assert not PLACEHOLDER_RE.findall(texto)
        assert "Nenhum periférico adicional vinculado." in texto
        assert "Não informado" in texto

        # A linha de especificações não pode ficar vazia (só o rótulo, sem conteúdo).
        for p in doc.paragraphs:
            if p.text.strip().lower().startswith("especifica"):
                assert p.text.strip() != "Especificações:", "Linha de especificações ficou vazia."

    def test_unidade_inexistente_produz_erro_explicativo(self, monkeypatch):
        ok, resultado, filename = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(revenda="Revalle Fictícia"), None
        )
        assert ok is False
        assert filename is None
        assert "Revalle Fictícia" in resultado
        assert "cadastr" in resultado.lower()

    def test_unidade_inativa_produz_erro_explicativo(self, monkeypatch):
        unidade_inativa = _unidade_padrao(nome="Revalle Inativa", is_active=0)
        ok, resultado, filename = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(revenda="Revalle Inativa"), unidade_inativa
        )
        assert ok is False
        assert filename is None
        assert "Revalle Inativa" in resultado
        assert "inativ" in resultado.lower()

    def test_cep_vazio_nao_deixa_cep_pendurado(self, monkeypatch):
        """Petrolina não tem CEP cadastrado — o termo não pode sair com
        'CEP ,' pendurado no meio da frase do bloco do EMPREGADOR."""
        unidade_sem_cep = _unidade_padrao(nome="Revalle Petrolina", cidade="Petrolina", uf="PE", cep="")
        ok, resultado, _ = _gerar_termo_emprestimo(
            monkeypatch, _item_padrao(revenda="Revalle Petrolina"), unidade_sem_cep
        )
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))
        texto = _texto_completo(doc)
        assert "CEP ," not in texto
        assert "CEP ." not in texto
        assert not PLACEHOLDER_RE.findall(texto)

    def test_cep_preenchido_aparece_no_texto(self, monkeypatch):
        unidade = _unidade_padrao(cep="48900-000")
        ok, resultado, _ = _gerar_termo_emprestimo(monkeypatch, _item_padrao(), unidade)
        assert ok, resultado
        doc = DocxDocument(io.BytesIO(resultado))
        assert "CEP 48900-000," in _texto_completo(doc)

    def test_data_hoje_extenso_em_portugues_sem_locale(self):
        data = inv_mod._data_por_extenso(datetime(2026, 7, 30))
        assert data == "30 de julho de 2026"


# ────────────────────────────────────────────────────────────────────────────
# Substituição run-aware (T1) — cobertura direta do utilitário
# ────────────────────────────────────────────────────────────────────────────

class TestSubstituicaoRunAware:
    def test_placeholder_dividido_entre_varios_runs_e_substituido_corretamente(self):
        """Monta um parágrafo com o texto fatiado de propósito entre runs (o Word
        faz isso rotineiramente ao salvar), como uma revisão manual do documento
        deixaria: rótulo em negrito num run, e o placeholder partido em dois
        runs sem formatação."""
        doc = DocxDocument()
        paragrafo = doc.add_paragraph()
        rotulo = paragrafo.add_run("Tipo de Equipamento:")
        rotulo.bold = True
        paragrafo.add_run(" ")
        parte1 = paragrafo.add_run("{{ti")
        parte2 = paragrafo.add_run("po}}")

        inv_mod._substituir_em_paragrafo(paragrafo, {"{{tipo}}": "Notebook"})

        assert paragrafo.text == "Tipo de Equipamento: Notebook"
        assert rotulo.bold is True
        assert rotulo.text == "Tipo de Equipamento:"
        # O run que iniciava o placeholder herda o valor; o restante é esvaziado,
        # nunca recriado (preserva a formatação original de cada run).
        assert parte1.text == "Notebook"
        assert parte2.text == ""

    def test_mesmo_placeholder_repetido_no_paragrafo_e_todo_substituido(self):
        doc = DocxDocument()
        paragrafo = doc.add_paragraph("{{cidade_uf}}, {{cidade_uf}}!")
        inv_mod._substituir_em_paragrafo(paragrafo, {"{{cidade_uf}}": "Juazeiro/BA"})
        assert paragrafo.text == "Juazeiro/BA, Juazeiro/BA!"

    def test_cabecalho_e_rodape_tambem_sao_substituidos(self):
        """T1: hoje só parágrafos e tabelas do corpo são varridos; um placeholder
        que caísse no cabeçalho/rodapé passaria intacto para o documento final."""
        doc = DocxDocument()
        secao = doc.sections[0]
        secao.header.is_linked_to_previous = False
        secao.footer.is_linked_to_previous = False
        secao.header.paragraphs[0].add_run("Unidade: {{cidade_uf}}")
        secao.footer.paragraphs[0].add_run("Rodapé - {{cidade_uf}}")

        inv_mod._substituir_placeholders_documento(doc, {"{{cidade_uf}}": "Juazeiro/BA"})

        assert "{{cidade_uf}}" not in secao.header.paragraphs[0].text
        assert "Juazeiro/BA" in secao.header.paragraphs[0].text
        assert "Juazeiro/BA" in secao.footer.paragraphs[0].text


# ────────────────────────────────────────────────────────────────────────────
# T4 — unidade nova sem modelo de devolução cadastrado
# ────────────────────────────────────────────────────────────────────────────

class TestMensagemDeErroDevolucaoSemModelo:
    def test_unidade_sem_entrada_em_termo_devolucao_modelos_gera_mensagem_clara(self):
        """Unidades criadas pela tela não têm entrada em TERMO_DEVOLUCAO_MODELOS
        (dicionário fixo por revenda) — a mensagem deve dizer explicitamente qual
        unidade está sem modelo de devolução e o que fazer, em vez do genérico
        'Modelo de termo de devolução não encontrado para {revenda}.'"""
        item = _item_padrao(
            status="Indisponível",
            revenda="Revalle Unidade Nova Sem Modelo",
            assigned_to="Fulano de Tal",
        )
        manager = _FakeManager(item, peripherals=[])
        ok, mensagem, filename = manager.generate_return_term_bytes(item["id"], "tester")
        assert ok is False
        assert filename is None
        assert "Revalle Unidade Nova Sem Modelo" in mensagem
        assert "TERMO_DEVOLUCAO_MODELOS" in mensagem
