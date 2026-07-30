"""
Camada de acesso ao banco de dados para o inventário.
Adaptado do original para uso no backend FastAPI:
- Importações corrigidas para pacote app.db.*
- generate_term() e generate_and_initiate_return() retornam bytes (DOCX em memória)
  em vez de salvar em disco — o router controla o storage.
- confirm_loan() e confirm_return() recebem storage_key (str) em vez de path local.
- remove() e replace_peripheral() recebem storage_key (str) em vez de path local.

Todos os métodos usam `database_mysql.get_cursor()`, que garante
commit/rollback/close sobre uma conexão obtida de um pool (em vez de abrir uma
conexão nova a cada chamada, como antes).
"""
import io
import os
import calendar
import logging
from datetime import datetime
import pymysql
from docx import Document

from app.db.database_mysql import get_cursor
from app.core.config import TERMO_MODELOS, TERMO_DEVOLUCAO_MODELOS
from app.db.utils import format_cpf, format_date
from app.db.unidade_db import get_unidade_manager

logger = logging.getLogger(__name__)

# ── Geração de termos (.docx) — utilitários compartilhados ─────────────────────
#
# Nomes dos meses em português para "{{data_hoje_extenso}}". Propositalmente NÃO
# usa o módulo `locale`: depende de configuração do sistema operacional (locale
# "pt_BR" instalado) e quebra silenciosamente em containers Linux mínimos, onde a
# formatação cai para o padrão "C" sem lançar exceção — um bug difícil de notar em
# produção e impossível de reproduzir localmente sem o mesmo container.
_MESES_PT_BR = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]


def _data_por_extenso(dt: datetime) -> str:
    """Formata uma data por extenso em português, ex.: '30 de julho de 2026'."""
    return f"{dt.day} de {_MESES_PT_BR[dt.month - 1]} de {dt.year}"


# Campos técnicos do equipamento incluídos em "{{especificacoes}}", na ordem em
# que devem aparecer, com o rótulo legível em português de cada um.
_CAMPOS_ESPECIFICACOES = [
    ("cpu", "CPU"),
    ("ram", "RAM"),
    ("storage", "Armazenamento"),
    ("sistema", "Sistema"),
    ("licenca", "Licença"),
    ("host", "Host"),
    ("ip", "IP"),
    ("mac", "MAC"),
    ("dominio", "Domínio"),
    ("anydesk", "AnyDesk"),
    ("potencia_nominal", "Potência Nominal"),
    ("autonomia_estimada", "Autonomia"),
    ("ip_snmp", "IP SNMP"),
    ("poe", "POE"),
    ("quantidade_portas", "Qtde. de Portas"),
    ("codigo_patrimonial", "Código Patrimonial"),
    ("local_instalacao", "Local de Instalação"),
]


def _montar_especificacoes(item: dict) -> str:
    """Junta os campos técnicos preenchidos do equipamento em uma única linha,
    para "{{especificacoes}}". Sem nenhum campo preenchido, devolve "Não
    informado" — nunca deixa o placeholder cru nem uma linha vazia no termo."""
    partes = [
        f"{rotulo}: {item[campo]}"
        for campo, rotulo in _CAMPOS_ESPECIFICACOES
        if item.get(campo) not in (None, "")
    ]
    return " | ".join(partes) if partes else "Não informado"


def _texto_runs(runs) -> str:
    return "".join(r.text for r in runs)


def _substituir_intervalo_em_runs(runs, inicio: int, fim: int, valor: str) -> None:
    """Escreve `valor` no lugar do intervalo [inicio, fim) do texto concatenado
    dos `runs`, preservando a formatação: o primeiro run atingido herda o texto
    novo (mantendo seu `bold`/`italic`/fonte), e os demais runs atingidos têm
    apenas a parte consumida removida (nunca são recriados)."""
    cursor = 0
    afetados = []  # (run, inicio_local, fim_local)
    for run in runs:
        run_inicio = cursor
        run_fim = cursor + len(run.text)
        if run_fim > inicio and run_inicio < fim:
            afetados.append((run, max(inicio, run_inicio) - run_inicio, min(fim, run_fim) - run_inicio))
        cursor = run_fim
        if cursor >= fim:
            break
    if not afetados:
        return
    primeiro_run, ini0, fim0 = afetados[0]
    texto = primeiro_run.text
    primeiro_run.text = texto[:ini0] + valor + texto[fim0:]
    for run, ini, fim_local in afetados[1:]:
        texto = run.text
        run.text = texto[:ini] + texto[fim_local:]


def _substituir_em_paragrafo(paragrafo, substituicoes: dict) -> None:
    """Substitui todas as ocorrências de cada chave de `substituicoes` no texto de
    um parágrafo, run-aware: localiza a chave no texto concatenado do parágrafo
    (o placeholder pode estar dividido entre vários runs — o Word faz isso
    rotineiramente ao salvar), escreve o valor no primeiro run do intervalo
    (herdando a formatação dele) e remove apenas a parte consumida dos runs
    seguintes. Diferente de `paragraph.text = ...`, nunca colapsa os runs nem
    descarta negrito/itálico/fonte."""
    runs = paragrafo.runs
    if not runs:
        return
    for chave, valor in substituicoes.items():
        valor = "" if valor is None else str(valor)
        while True:
            texto_atual = _texto_runs(runs)
            pos = texto_atual.find(chave)
            if pos == -1:
                break
            _substituir_intervalo_em_runs(runs, pos, pos + len(chave), valor)


def _paragrafos_de(container):
    """Gera todos os parágrafos de um container do python-docx (documento,
    célula de tabela, cabeçalho ou rodapé), incluindo os de tabelas aninhadas,
    recursivamente."""
    for paragrafo in container.paragraphs:
        yield paragrafo
    for tabela in container.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                yield from _paragrafos_de(celula)


def _substituir_placeholders_documento(doc: Document, substituicoes: dict) -> None:
    """Aplica `_substituir_em_paragrafo` em TODO o documento: parágrafos e
    tabelas do corpo, e cabeçalhos/rodapés de todas as seções (incluindo
    variantes de primeira página e página par/ímpar, quando a seção as usa) —
    um placeholder que caísse só no cabeçalho passaria intacto para o
    documento final se apenas o corpo fosse varrido."""
    containers = [doc]
    for secao in doc.sections:
        containers.extend([
            secao.header, secao.footer,
            secao.first_page_header, secao.first_page_footer,
            secao.even_page_header, secao.even_page_footer,
        ])
    for container in containers:
        for paragrafo in _paragrafos_de(container):
            _substituir_em_paragrafo(paragrafo, substituicoes)


def _texto_perifericos(linked_peripherals, texto_vazio: str) -> str:
    if not linked_peripherals:
        return texto_vazio
    return "\n".join(
        f"- {p['tipo']}: {p.get('brand','')} {p.get('model','')} (S/N: {p.get('identificador') or 'N/A'})"
        for p in linked_peripherals
    )


class InventoryDBManager:
    def __init__(self):
        self._create_tables()

    def _create_tables(self):
        with get_cursor(dict_cursor=False) as cur:
            cur.execute("""
            CREATE TABLE IF NOT EXISTS items (
                id INT AUTO_INCREMENT PRIMARY KEY,
                tipo VARCHAR(50), brand VARCHAR(100), model VARCHAR(100), identificador VARCHAR(100),
                nota_fiscal VARCHAR(50),
                status ENUM('Disponível','Indisponível','Pendente','Pendente Devolução') DEFAULT 'Disponível',
                assigned_to VARCHAR(100), cpf VARCHAR(20), revenda VARCHAR(100),
                dominio VARCHAR(50), host VARCHAR(100), endereco_fisico VARCHAR(150),
                cpu VARCHAR(100), ram VARCHAR(50), storage VARCHAR(50),
                sistema VARCHAR(100), licenca VARCHAR(100), anydesk VARCHAR(50),
                setor VARCHAR(100), ip VARCHAR(50), mac VARCHAR(50),
                fornecedor VARCHAR(150),
                potencia_nominal VARCHAR(50), autonomia_estimada VARCHAR(100), ip_snmp VARCHAR(50),
                codigo_patrimonial VARCHAR(100), responsavel VARCHAR(100), local_instalacao VARCHAR(150),
                poe ENUM('Sim','Não'), quantidade_portas VARCHAR(10),
                date_registered DATETIME NOT NULL, date_issued DATETIME,
                is_active TINYINT(1) DEFAULT 1,
                pessoa_juridica TINYINT(1) DEFAULT 0
            )
            """)
            cur.execute("""
            CREATE TABLE IF NOT EXISTS history (
                id INT AUTO_INCREMENT PRIMARY KEY,
                item_id INT, peripheral_id INT,
                operador VARCHAR(100), usuario VARCHAR(100), cpf VARCHAR(20),
                cargo VARCHAR(100), center_cost VARCHAR(100), setor VARCHAR(100),
                fornecedor VARCHAR(150), revenda VARCHAR(100),
                data_operacao DATETIME,
                operation ENUM(
                    'Cadastro','Empréstimo','Devolução','Edição','Exclusão','Estorno',
                    'Confirmação Empréstimo','Confirmação Devolução',
                    'Cadastro Periférico','Vínculo Periférico','Desvínculo Periférico','Substituição Periférico'
                ) DEFAULT 'Cadastro',
                is_reversed TINYINT(1) DEFAULT 0,
                details VARCHAR(255),
                tipo VARCHAR(50), brand VARCHAR(100), model VARCHAR(100), identificador VARCHAR(100),
                nota_fiscal VARCHAR(50), poe ENUM('Sim','Não'), quantidade_portas VARCHAR(10),
                operacao_anexo VARCHAR(255) NULL, termo_assinado_anexo VARCHAR(255) NULL,
                pessoa_juridica TINYINT(1) DEFAULT 0,
                FOREIGN KEY(item_id) REFERENCES items(id) ON DELETE SET NULL
            )
            """)

            # Migração leve para bancos onde a tabela 'items'/'history' já existia
            # antes deste campo ser criado: CREATE TABLE IF NOT EXISTS não altera
            # uma tabela existente, então a coluna não apareceria sozinha num banco
            # com dados reais (como o de produção). Idempotente — não corre risco
            # de tentar adicionar a coluna duas vezes.
            for tabela in ("items", "history"):
                cur.execute(
                    "SELECT COUNT(*) FROM information_schema.columns "
                    "WHERE table_schema = DATABASE() AND table_name = %s "
                    "AND column_name = 'pessoa_juridica'",
                    (tabela,),
                )
                if cur.fetchone()[0] == 0:
                    cur.execute(f"ALTER TABLE {tabela} ADD COLUMN pessoa_juridica TINYINT(1) DEFAULT 0")
            cur.execute("""
            CREATE TABLE IF NOT EXISTS peripherals (
                id INT AUTO_INCREMENT PRIMARY KEY,
                tipo VARCHAR(50) NOT NULL, brand VARCHAR(100), model VARCHAR(100),
                identificador VARCHAR(100) UNIQUE,
                status ENUM('Disponível','Em Uso','Substituido') DEFAULT 'Disponível',
                motivo_substituicao VARCHAR(255),
                date_registered DATETIME NOT NULL,
                is_active TINYINT(1) DEFAULT 1
            )
            """)
            cur.execute("""
            CREATE TABLE IF NOT EXISTS equipment_peripherals (
                id INT AUTO_INCREMENT PRIMARY KEY,
                equipment_id INT NOT NULL, peripheral_id INT NOT NULL,
                FOREIGN KEY(equipment_id) REFERENCES items(id) ON DELETE CASCADE,
                FOREIGN KEY(peripheral_id) REFERENCES peripherals(id) ON DELETE CASCADE,
                UNIQUE(equipment_id, peripheral_id)
            )
            """)

    # ── Items ──────────────────────────────────────────────────────────────────

    def add_item(self, item_data: dict, logged_user: str):
        keys = ", ".join(item_data.keys())
        placeholders = ", ".join(["%s"] * len(item_data))
        sql = f"INSERT INTO items ({keys}) VALUES ({placeholders})"
        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(sql, list(item_data.values()))
                item_id = cur.lastrowid
                cur.execute(
                    "INSERT INTO history (item_id, operador, data_operacao, operation) VALUES (%s,%s,%s,'Cadastro')",
                    (item_id, logged_user, datetime.now()),
                )
            return True, item_id
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def update_item(self, item_id: int, item_data: dict, logged_user: str):
        sets = ", ".join([f"{k}=%s" for k in item_data.keys()])
        values = list(item_data.values()) + [item_id]
        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(f"UPDATE items SET {sets} WHERE id=%s", values)
                cur.execute(
                    "INSERT INTO history (item_id, operador, data_operacao, operation) VALUES (%s,%s,%s,'Edição')",
                    (item_id, logged_user, datetime.now()),
                )
            return True, f"Item {item_id} atualizado."
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def remove(self, item_id: int, logged_user: str, reason: str, storage_key: str = None):
        """Remove (soft-delete) um item. storage_key é a chave do anexo no storage backend."""
        item = self.find(item_id)
        if not item:
            return False, "ID não encontrado."
        if item["status"] != "Disponível":
            return False, "Não é possível remover produto emprestado."

        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute("UPDATE items SET is_active=0 WHERE id=%s", (item_id,))
                cur.execute(
                    """INSERT INTO history
                    (item_id, operador, data_operacao, operation, details,
                     tipo, brand, model, identificador, nota_fiscal, fornecedor,
                     operacao_anexo, poe, quantidade_portas)
                    VALUES (%s,%s,%s,'Exclusão',%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                    (
                        item_id, logged_user, datetime.now(), reason,
                        item.get("tipo"), item.get("brand"), item.get("model"),
                        item.get("identificador"), item.get("nota_fiscal"),
                        item.get("fornecedor"), storage_key,
                        item.get("poe"), item.get("quantidade_portas"),
                    ),
                )
            return True, f"Aparelho {item_id} removido do estoque."
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def list_items(
        self,
        *,
        tipo: str | None = None,
        status: str | None = None,
        revenda: str | None = None,
        search: str | None = None,
        limit: int | None = None,
        offset: int = 0,
    ) -> tuple[list[dict], int]:
        """Lista itens ativos aplicando filtros e paginação em SQL.

        Retorna (linhas, total_sem_paginacao). `limit=None` retorna tudo
        (mantém compatibilidade com quem não pagina). `search` varre brand,
        model, identificador, assigned_to e host via LIKE (sempre com
        placeholders `%s` — nunca f-string, para não abrir brecha de SQL
        injection).
        """
        where = ["i.is_active = 1"]
        params: list = []
        if tipo:
            where.append("i.tipo = %s")
            params.append(tipo)
        if status:
            where.append("i.status = %s")
            params.append(status)
        if revenda:
            where.append("i.revenda = %s")
            params.append(revenda)
        if search:
            like = f"%{search}%"
            where.append(
                "(i.brand LIKE %s OR i.model LIKE %s OR i.identificador LIKE %s "
                "OR i.assigned_to LIKE %s OR i.host LIKE %s)"
            )
            params.extend([like, like, like, like, like])
        where_sql = " AND ".join(where)

        with get_cursor(commit=False) as cur:
            cur.execute(f"SELECT COUNT(*) as total FROM items i WHERE {where_sql}", params)
            total = cur.fetchone()["total"]

            # GROUP BY (por causa do peripheral_count) impede reaproveitar o
            # COUNT(*) direto na mesma query — por isso a contagem acima é
            # feita à parte, sem o LEFT JOIN/GROUP BY.
            query = f"""
                SELECT i.*, COUNT(ep.peripheral_id) as peripheral_count
                FROM items i
                LEFT JOIN equipment_peripherals ep ON i.id = ep.equipment_id
                WHERE {where_sql}
                GROUP BY i.id
                ORDER BY i.id DESC
            """
            query_params = list(params)
            if limit is not None:
                query += " LIMIT %s OFFSET %s"
                query_params.extend([limit, offset])
            cur.execute(query, query_params)
            rows = cur.fetchall()
        return rows, total

    def find(self, item_id: int):
        with get_cursor(commit=False) as cur:
            cur.execute("SELECT * FROM items WHERE id=%s AND is_active=1", (item_id,))
            return cur.fetchone()

    # ── Peripherals ────────────────────────────────────────────────────────────

    def add_peripheral(self, data: dict, logged_user: str):
        try:
            data["date_registered"] = datetime.now()
            keys = ", ".join(data.keys())
            placeholders = ", ".join(["%s"] * len(data))
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(f"INSERT INTO peripherals ({keys}) VALUES ({placeholders})", list(data.values()))
                p_id = cur.lastrowid
                cur.execute(
                    """INSERT INTO history (peripheral_id, operador, data_operacao, operation, tipo, brand, model, identificador)
                    VALUES (%s,%s,%s,'Cadastro Periférico',%s,%s,%s,%s)""",
                    (p_id, logged_user, datetime.now(), data.get("tipo"), data.get("brand"), data.get("model"), data.get("identificador")),
                )
            return True, f"Periférico '{data.get('tipo')}' cadastrado com ID {p_id}."
        except pymysql.MySQLError as e:
            if e.args[0] == 1062:
                return False, "Já existe um periférico com este Identificador (Nº de Série)."
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def list_peripherals(self, status_filter="", type_filter="", include_inactive=False):
        where = []
        params = []
        if not include_inactive:
            where.append("is_active=1")
        if status_filter:
            where.append("status=%s")
            params.append(status_filter)
        if type_filter:
            where.append("tipo=%s")
            params.append(type_filter)
        query = "SELECT * FROM peripherals"
        if where:
            query += " WHERE " + " AND ".join(where)
        with get_cursor(commit=False) as cur:
            cur.execute(query, params)
            return cur.fetchall()

    def deactivate_peripheral(self, peripheral_id: int, logged_user: str) -> tuple[bool, str]:
        """Soft-delete de periférico. Falha se estiver 'Em Uso' ou vinculado a
        algum equipamento (mesma regra de negócio do soft-delete de item)."""
        with get_cursor(commit=False) as cur:
            cur.execute("SELECT * FROM peripherals WHERE id=%s AND is_active=1", (peripheral_id,))
            peripheral = cur.fetchone()
            if not peripheral:
                return False, "Periférico não encontrado."
            if peripheral["status"] == "Em Uso":
                return False, "Não é possível remover periférico em uso."
            cur.execute(
                "SELECT COUNT(*) as total FROM equipment_peripherals WHERE peripheral_id=%s",
                (peripheral_id,),
            )
            if cur.fetchone()["total"] > 0:
                return False, "Não é possível remover periférico vinculado a um equipamento."

        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute("UPDATE peripherals SET is_active=0 WHERE id=%s", (peripheral_id,))
                cur.execute(
                    """INSERT INTO history (peripheral_id, operador, data_operacao, operation, tipo, brand, model, identificador)
                    VALUES (%s,%s,%s,'Exclusão',%s,%s,%s,%s)""",
                    (
                        peripheral_id, logged_user, datetime.now(),
                        peripheral.get("tipo"), peripheral.get("brand"),
                        peripheral.get("model"), peripheral.get("identificador"),
                    ),
                )
            return True, f"Periférico '{peripheral.get('tipo')}' (ID {peripheral_id}) removido do estoque."
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def list_peripherals_for_equipment(self, equipment_id: int):
        with get_cursor(commit=False) as cur:
            cur.execute(
                """SELECT p.*, ep.id as link_id
                FROM peripherals p
                JOIN equipment_peripherals ep ON p.id = ep.peripheral_id
                WHERE ep.equipment_id = %s""",
                (equipment_id,),
            )
            return cur.fetchall()

    def link_peripheral_to_equipment(self, equipment_id: int, peripheral_id: int, logged_user: str):
        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(
                    "INSERT INTO equipment_peripherals (equipment_id, peripheral_id) VALUES (%s,%s)",
                    (equipment_id, peripheral_id),
                )
                cur.execute("UPDATE peripherals SET status='Em Uso' WHERE id=%s", (peripheral_id,))
                cur.execute(
                    """INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation)
                    VALUES (%s,%s,%s,%s,'Vínculo Periférico')""",
                    (equipment_id, peripheral_id, logged_user, datetime.now()),
                )
            return True, "Periférico vinculado com sucesso."
        except pymysql.MySQLError as e:
            # equipment_peripherals tem UNIQUE(equipment_id, peripheral_id), então
            # vincular o mesmo par duas vezes cai em 1062. É erro de uso, não de
            # sistema: merece mensagem específica, como já acontece em add_peripheral.
            if e.args and e.args[0] == 1062:
                return False, "Este periférico já está vinculado a este equipamento."
            logger.exception("Erro ao vincular periférico")
            return False, "Erro ao vincular periférico."

    def unlink_peripheral_from_equipment(self, link_id: int, logged_user: str):
        try:
            with get_cursor() as cur:
                cur.execute("SELECT equipment_id, peripheral_id FROM equipment_peripherals WHERE id=%s", (link_id,))
                ids = cur.fetchone()
                if not ids:
                    return False, "Vínculo não encontrado."
                cur.execute("DELETE FROM equipment_peripherals WHERE id=%s", (link_id,))
                cur.execute("UPDATE peripherals SET status='Disponível' WHERE id=%s", (ids["peripheral_id"],))
                cur.execute(
                    """INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation)
                    VALUES (%s,%s,%s,%s,'Desvínculo Periférico')""",
                    (ids["equipment_id"], ids["peripheral_id"], logged_user, datetime.now()),
                )
            return True, "Periférico desvinculado com sucesso."
        except pymysql.MySQLError:
            logger.exception("Erro ao desvincular periférico")
            return False, "Erro ao desvincular periférico."

    def replace_peripheral(
        self,
        equipment_id: int,
        old_peripheral_id: int,
        new_peripheral_id: int,
        reason: str,
        logged_user: str,
        storage_key: str = None,
    ):
        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(
                    "UPDATE peripherals SET status='Substituido', is_active=0, motivo_substituicao=%s WHERE id=%s",
                    (reason, old_peripheral_id),
                )
                cur.execute(
                    "DELETE FROM equipment_peripherals WHERE equipment_id=%s AND peripheral_id=%s",
                    (equipment_id, old_peripheral_id),
                )
                cur.execute(
                    "INSERT INTO equipment_peripherals (equipment_id, peripheral_id) VALUES (%s,%s)",
                    (equipment_id, new_peripheral_id),
                )
                cur.execute("UPDATE peripherals SET status='Em Uso' WHERE id=%s", (new_peripheral_id,))
                details_log = f"Substituído periférico ID {old_peripheral_id} por ID {new_peripheral_id}. Motivo: {reason}"
                cur.execute(
                    """INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation, details, operacao_anexo)
                    VALUES (%s,%s,%s,%s,'Substituição Periférico',%s,%s)""",
                    (equipment_id, old_peripheral_id, logged_user, datetime.now(), details_log, storage_key),
                )
            return True, "Substituição realizada com sucesso."
        except pymysql.MySQLError:
            logger.exception("Erro ao substituir periférico")
            return False, "Erro ao substituir periférico."

    # ── Loan workflow ──────────────────────────────────────────────────────────

    def issue(
        self, pid, user, cpf, center_cost, cargo, setor, revenda, date_issue, logged_user: str,
        pessoa_juridica: bool = False,
    ):
        item = self.find(pid)
        if not item:
            return False, "Item não encontrado."
        if item["status"] != "Disponível":
            return False, "Este item não está disponível para empréstimo."
        try:
            dt_issue = datetime.strptime(date_issue, "%d/%m/%Y")
            dt_issue = dt_issue.replace(
                hour=datetime.now().hour,
                minute=datetime.now().minute,
                second=datetime.now().second,
            )
        except ValueError:
            return False, "Data de empréstimo inválida (use dd/mm/aaaa)."
        if dt_issue.date() > datetime.now().date():
            return False, "A data de empréstimo não pode ser no futuro."
        if item.get("date_registered") and dt_issue.date() < item["date_registered"].date():
            return False, f"Data de empréstimo não pode ser anterior ao cadastro ({format_date(str(item['date_registered']))})."

        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute(
                    "UPDATE items SET status='Pendente', assigned_to=%s, cpf=%s, date_issued=%s, revenda=%s, pessoa_juridica=%s WHERE id=%s",
                    (user, cpf, dt_issue, revenda, pessoa_juridica, pid),
                )
                # pessoa_juridica também vai para o histórico (não só para items):
                # é o que permite reverse_history_entry() restaurar o valor correto
                # se uma Devolução for estornada depois (ver esse método mais abaixo).
                cur.execute(
                    """INSERT INTO history (item_id, operador, usuario, cpf, cargo, center_cost, setor, revenda, fornecedor, data_operacao, operation, pessoa_juridica)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'Empréstimo',%s)""",
                    (pid, logged_user, user, cpf, cargo, center_cost, setor, revenda, item.get("fornecedor"), dt_issue, pessoa_juridica),
                )
            return True, f"Empréstimo do item {pid} para {user} iniciado. Status: Pendente."
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente."

    def confirm_loan(self, item_id: int, logged_user: str, storage_key: str):
        """Confirma empréstimo. storage_key é a chave do PDF assinado no storage backend."""
        item = self.find(item_id)
        if not item or item["status"] != "Pendente":
            return False, "Apenas itens com status 'Pendente' podem ser confirmados."

        try:
            with get_cursor() as cur:
                cur.execute(
                    """SELECT id, usuario, cpf, cargo, center_cost, revenda, pessoa_juridica FROM history
                    WHERE item_id=%s AND operation='Empréstimo' AND is_reversed=0
                    ORDER BY data_operacao DESC, id DESC LIMIT 1""",
                    (item_id,),
                )
                last_loan = cur.fetchone()
                if not last_loan:
                    raise Exception("Registro de empréstimo original não encontrado no histórico.")

                cur.execute("UPDATE items SET status='Indisponível' WHERE id=%s", (item_id,))

                # Antes chamava self.list_peripherals_for_equipment(item_id), o que abria
                # uma segunda conexão em paralelo à desta transação. Agora usa o mesmo cursor.
                cur.execute(
                    """SELECT p.id FROM peripherals p
                    JOIN equipment_peripherals ep ON p.id = ep.peripheral_id
                    WHERE ep.equipment_id = %s""",
                    (item_id,),
                )
                peripheral_ids = [p["id"] for p in cur.fetchall()]
                for p_id in peripheral_ids:
                    cur.execute("UPDATE peripherals SET status='Em Uso' WHERE id=%s", (p_id,))

                cur.execute(
                    """INSERT INTO history (item_id, operador, usuario, cpf, cargo, center_cost, revenda, data_operacao, operation, termo_assinado_anexo, pessoa_juridica)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'Confirmação Empréstimo',%s,%s)""",
                    (
                        item_id, logged_user,
                        last_loan.get("usuario"), last_loan.get("cpf"),
                        last_loan.get("cargo"), last_loan.get("center_cost"),
                        last_loan.get("revenda"), datetime.now(), storage_key,
                        last_loan.get("pessoa_juridica") or 0,
                    ),
                )
            return True, f"Empréstimo do item {item_id} confirmado."
        except Exception:
            logger.exception("Erro ao confirmar empréstimo")
            return False, "Erro ao confirmar empréstimo."

    def get_active_signed_term_key(self, item_id: int):
        """Retorna a storage_key do termo de empréstimo assinado vigente (não
        estornado) para o item, ou None se não houver confirmação ativa.

        Filtra is_reversed=0: sem esse filtro, um empréstimo já estornado
        (a confirmação foi desfeita) ainda poderia "emprestar" o termo
        assinado de uma operação que não vale mais.
        """
        with get_cursor(commit=False) as cur:
            cur.execute(
                """SELECT termo_assinado_anexo FROM history
                WHERE item_id=%s AND operation='Confirmação Empréstimo' AND is_reversed=0
                ORDER BY data_operacao DESC, id DESC LIMIT 1""",
                (item_id,),
            )
            row = cur.fetchone()
        return row["termo_assinado_anexo"] if row else None

    # ── Return workflow ────────────────────────────────────────────────────────

    def generate_return_term_bytes(self, item_id: int, logged_user: str):
        """
        Gera o termo de devolução em memória e inicia o processo (status → Pendente Devolução).
        Retorna (True, bytes_do_docx, filename) ou (False, msg_erro, None).
        """
        item = self.find(item_id)
        if not item or item["status"] != "Indisponível":
            return False, "Apenas itens com status 'Indisponível' podem ter termo de devolução gerado.", None

        user = item.get("assigned_to")
        if not user:
            return False, "Não foi possível encontrar o usuário associado.", None

        revenda = item.get("revenda")
        # Agora que unidades podem ser cadastradas pela tela, uma unidade nova não
        # terá entrada em TERMO_DEVOLUCAO_MODELOS (dicionário fixo por revenda) — sem
        # essa checagem em duas etapas, a devolução falharia com a mensagem genérica
        # "Modelo de termo de devolução não encontrado para {revenda}", sem indicar o
        # que fazer. Não inventamos um modelo de devolução para a unidade: apenas
        # explicamos que é preciso providenciar e cadastrar o .docx correspondente.
        if revenda not in TERMO_DEVOLUCAO_MODELOS:
            return False, (
                f"Não há modelo de termo de devolução cadastrado para a unidade '{revenda}'. "
                "Esta é provavelmente uma unidade nova, criada pela tela de unidades: "
                "providencie o arquivo .docx do termo de devolução dessa unidade e registre-o "
                "em TERMO_DEVOLUCAO_MODELOS (backend/app/core/config.py) antes de devolver "
                "equipamentos vinculados a ela."
            ), None
        modelo_path = TERMO_DEVOLUCAO_MODELOS[revenda]
        if not os.path.exists(modelo_path):
            return False, (
                f"O modelo de termo de devolução cadastrado para '{revenda}' não foi encontrado "
                f"em {modelo_path}. Verifique se o arquivo .docx está presente em backend/modelos/devolucao/."
            ), None

        linked_peripherals = self.list_peripherals_for_equipment(item_id)
        peripherals_text = _texto_perifericos(linked_peripherals, "Nenhum periférico adicional devolvido.")

        detalhes_parts = [f"{item.get('tipo','')} {item.get('brand','')} {item.get('model','')}".strip()]
        for key, label in {"identificador": "S/N", "cpu": "CPU", "ram": "RAM",
                           "storage": "Armazenamento", "setor": "Setor", "ip": "IP", "mac": "MAC"}.items():
            if item.get(key):
                detalhes_parts.append(f"{label}: {item[key]}")
        detalhes_finais = " - ".join(detalhes_parts)

        substituicoes = {
            "{{nome}}": user or "",
            "{{data_hoje}}": datetime.now().strftime("%d/%m/%Y"),
            "{{cpf}}": format_cpf(item.get("cpf", "")),
            "{{data_emprestimo}}": format_date(item.get("date_issued", "")),
            "{{tipo}}": item.get("tipo", ""),
            "{{detalhes_equipamento}}": detalhes_finais,
            "{{perifericos}}": peripherals_text,
        }
        substituicoes = {k: (v if v is not None else "") for k, v in substituicoes.items()}

        try:
            doc = Document(modelo_path)
            _substituir_placeholders_documento(doc, substituicoes)
            buf = io.BytesIO()
            doc.save(buf)
            doc_bytes = buf.getvalue()
        except Exception:
            logger.exception("Erro ao gerar documento")
            return False, "Erro ao gerar o documento.", None

        # Atualiza status e registra no histórico
        try:
            with get_cursor(dict_cursor=False) as cur:
                cur.execute("UPDATE items SET status='Pendente Devolução' WHERE id=%s", (item_id,))
                cur.execute(
                    "INSERT INTO history (item_id, operador, data_operacao, operation) VALUES (%s,%s,%s,'Devolução')",
                    (item_id, logged_user, datetime.now()),
                )
        except pymysql.MySQLError:
            logger.exception("Erro de banco de dados")
            return False, "Erro ao processar a operação. Tente novamente.", None

        safe_user = user.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"termo_devolucao_{item_id}_{safe_user}_{revenda}_{timestamp}.docx"
        return True, doc_bytes, filename

    def confirm_return(self, item_id: int, logged_user: str, storage_key: str):
        """Confirma devolução. storage_key é a chave do PDF assinado no storage backend."""
        item = self.find(item_id)
        if not item or item["status"] != "Pendente Devolução":
            return False, "Apenas itens com status 'Pendente Devolução' podem ser confirmados."

        try:
            with get_cursor() as cur:
                cur.execute(
                    "UPDATE items SET status='Disponível', assigned_to=NULL, cpf=NULL, date_issued=NULL, pessoa_juridica=0 WHERE id=%s",
                    (item_id,),
                )

                # Antes chamava self.list_peripherals_for_equipment(item_id) (2a conexão);
                # agora usa o mesmo cursor da transação.
                cur.execute(
                    """SELECT p.id FROM peripherals p
                    JOIN equipment_peripherals ep ON p.id = ep.peripheral_id
                    WHERE ep.equipment_id = %s""",
                    (item_id,),
                )
                peripheral_ids = [p["id"] for p in cur.fetchall()]
                for p_id in peripheral_ids:
                    cur.execute("UPDATE peripherals SET status='Disponível' WHERE id=%s", (p_id,))

                # Bug corrigido: antes o vínculo em equipment_peripherals nunca era
                # removido na devolução — o periférico ficava "Disponível" e, ao
                # mesmo tempo, vinculado ao equipamento (podendo ser vinculado a um
                # segundo equipamento). Remove os vínculos e registra o desvínculo
                # no histórico para manter a rastreabilidade.
                if peripheral_ids:
                    cur.execute(
                        "DELETE FROM equipment_peripherals WHERE equipment_id=%s",
                        (item_id,),
                    )
                    for p_id in peripheral_ids:
                        cur.execute(
                            """INSERT INTO history (item_id, peripheral_id, operador, data_operacao, operation)
                            VALUES (%s,%s,%s,%s,'Desvínculo Periférico')""",
                            (item_id, p_id, logged_user, datetime.now()),
                        )

                cur.execute(
                    """INSERT INTO history (item_id, operador, data_operacao, operation, termo_assinado_anexo)
                    VALUES (%s,%s,%s,'Confirmação Devolução',%s)""",
                    (item_id, logged_user, datetime.now(), storage_key),
                )
            return True, f"Devolução do item {item_id} confirmada."
        except Exception:
            logger.exception("Erro ao confirmar devolução")
            return False, "Erro ao confirmar devolução."

    # ── Loan term generation ───────────────────────────────────────────────────

    def generate_loan_term_bytes(self, item_id: int):
        """
        Gera o termo de responsabilidade em memória, a partir do template único
        (TERMO_MODELOS, ver app/core/config.py) preenchido com os dados da unidade
        (bloco do EMPREGADOR) e do equipamento/empréstimo.
        Retorna (True, bytes_do_docx, filename) ou (False, msg_erro, None).
        """
        item = self.find(item_id)
        if not item:
            return False, "Equipamento não encontrado.", None
        if item["status"] != "Pendente":
            return False, "Este equipamento não está pendente de empréstimo.", None

        revenda = item.get("revenda")

        # O bloco do EMPREGADOR é preenchido com os dados da unidade cadastrada no
        # banco — nunca deixamos o termo (documento jurídico) sair com esse bloco
        # vazio ou parcial. Se a unidade não existir ou estiver inativa, a geração
        # é recusada com uma mensagem que diz explicitamente qual unidade precisa
        # ser cadastrada/ativada, em vez de produzir um termo incompleto.
        unidade = get_unidade_manager().get_unidade_por_nome(revenda)
        if not unidade:
            return False, (
                f"A unidade '{revenda}' não está cadastrada. Cadastre a unidade "
                "(com razão social, CNPJ e endereço) antes de gerar o termo de "
                "responsabilidade."
            ), None
        if not unidade.get("is_active"):
            return False, (
                f"A unidade '{revenda}' está inativa. Reative a unidade antes de "
                "gerar o termo de responsabilidade."
            ), None

        modelo_path = TERMO_MODELOS
        if not os.path.exists(modelo_path):
            return False, f"Modelo de termo de responsabilidade não encontrado em {modelo_path}.", None

        linked_peripherals = self.list_peripherals_for_equipment(item_id)
        peripherals_text = _texto_perifericos(linked_peripherals, "Nenhum periférico adicional vinculado.")

        cidade = (unidade.get("cidade") or "").strip()
        uf = (unidade.get("uf") or "").strip()
        cidade_uf = f"{cidade}/{uf}" if cidade and uf else cidade

        # CEP pode estar vazio (ex.: Petrolina não tem CEP cadastrado). O trecho
        # do template é "..., CEP {{cep}}, inscrita ..." — se substituíssemos
        # apenas "{{cep}}" por "", o termo sairia com "CEP , inscrita" pendurado.
        # Por isso a chave de substituição é o trecho inteiro "CEP {{cep}}, ":
        # com CEP presente vira "CEP <cep>, "; sem CEP, o segmento inteiro
        # desaparece e a frase segue direto de "{{cidade_uf}}," para "inscrita".
        cep = (unidade.get("cep") or "").strip()
        cep_segmento = f"CEP {cep}, " if cep else ""

        user = item.get("assigned_to", "")
        marca_modelo = f"{item.get('brand','')} {item.get('model','')}".strip()

        substituicoes = {
            "CEP {{cep}}, ": cep_segmento,
            "{{razao_social}}": unidade.get("razao_social") or "",
            "{{endereco}}": unidade.get("endereco") or "",
            "{{cidade_uf}}": cidade_uf,
            "{{cep}}": cep,  # rede de segurança: nunca deixa o token cru se o segmento acima não bater
            "{{cnpj}}": unidade.get("cnpj") or "",
            "{{nome}}": user or "",
            # Corrige um defeito do documento original: o parágrafo do empregado
            # dizia sempre "pessoa jurídica de direito privado", mas o campo ao
            # lado é CPF — empregado é pessoa física na esmagadora maioria dos
            # casos. O padrão agora é "pessoa física"; o toggle pessoa_juridica
            # (desmarcado por padrão no formulário de empréstimo) permite marcar
            # o caso em que o recebedor é de fato representante de pessoa jurídica.
            "{{tipo_pessoa}}": "pessoa jurídica de direito privado" if item.get("pessoa_juridica") else "pessoa física",
            "{{cpf}}": format_cpf(item.get("cpf", "")),
            "{{tipo}}": item.get("tipo", "") or "",
            "{{marca_modelo}}": marca_modelo,
            "{{identificador}}": item.get("identificador", "") or "",
            "{{perifericos}}": peripherals_text,
            "{{especificacoes}}": _montar_especificacoes(item),
            "{{data_hoje_extenso}}": _data_por_extenso(datetime.now()),
        }
        substituicoes = {k: (v if v is not None else "") for k, v in substituicoes.items()}

        try:
            doc = Document(modelo_path)
            _substituir_placeholders_documento(doc, substituicoes)
            buf = io.BytesIO()
            doc.save(buf)
            doc_bytes = buf.getvalue()
        except Exception:
            logger.exception("Erro ao gerar documento")
            return False, "Erro ao gerar o documento.", None

        safe_user = str(user).replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"termo_{item_id}_{safe_user}_{revenda}_{timestamp}.docx"
        return True, doc_bytes, filename

    # ── History ────────────────────────────────────────────────────────────────

    def list_history(
        self, *, search: str | None = None, limit: int | None = None, offset: int = 0
    ) -> tuple[list[dict], int]:
        """Retorna (linhas, total_sem_paginacao).

        A busca é feita no banco, e não na página já carregada: com paginação
        server-side, filtrar no cliente só varreria as linhas visíveis e daria a
        impressão falsa de que o histórico inteiro foi pesquisado.
        """
        # As colunas exibidas vêm de COALESCE entre history, items e peripherals,
        # e o MySQL não aceita alias de SELECT no WHERE — por isso as expressões
        # são repetidas aqui.
        campos_busca = [
            "h.operador",
            "h.usuario",
            "h.operation",
            "h.revenda",
            "COALESCE(i.tipo, p.tipo, h.tipo)",
            "COALESCE(i.brand, p.brand, h.brand)",
            "COALESCE(i.model, p.model, h.model)",
            "COALESCE(i.identificador, p.identificador, h.identificador)",
        ]

        condicoes = ["h.is_reversed = 0"]
        filtro_params: list = []
        if search and search.strip():
            termo = f"%{search.strip()}%"
            condicoes.append("(" + " OR ".join(f"{c} LIKE %s" for c in campos_busca) + ")")
            filtro_params.extend([termo] * len(campos_busca))

        origem = f"""
            FROM history h
            LEFT JOIN items i ON i.id = h.item_id
            LEFT JOIN peripherals p ON p.id = h.peripheral_id
            WHERE {" AND ".join(condicoes)}
        """

        with get_cursor(commit=False) as cur:
            cur.execute(f"SELECT COUNT(*) as total {origem}", filtro_params)
            total = cur.fetchone()["total"]

            query = f"""
                SELECT h.id, h.item_id, h.operador, h.peripheral_id, h.details,
                    COALESCE(i.tipo, p.tipo, h.tipo) as tipo,
                    COALESCE(i.brand, p.brand, h.brand) as marca,
                    COALESCE(i.model, p.model, h.model) as modelo,
                    COALESCE(i.identificador, p.identificador, h.identificador) as identificador,
                    COALESCE(i.nota_fiscal, h.nota_fiscal) as nota_fiscal,
                    COALESCE(i.fornecedor, h.fornecedor) as fornecedor,
                    h.usuario, h.cpf, h.cargo, h.center_cost, h.setor, h.revenda,
                    h.operacao_anexo, h.termo_assinado_anexo, h.pessoa_juridica,
                    h.data_operacao, h.operation
                {origem}
                ORDER BY h.data_operacao DESC, h.id DESC
            """
            params = list(filtro_params)
            if limit is not None:
                query += " LIMIT %s OFFSET %s"
                params.extend([limit, offset])
            cur.execute(query, params)
            rows = cur.fetchall()
        return rows, total

    def generate_monthly_report(self, ano, mes):
        sql = """
            (SELECT h.id AS history_id, h.item_id, h.peripheral_id, h.operador,
                h.usuario, h.cpf, h.cargo, h.center_cost, h.setor,
                COALESCE(i.fornecedor, h.fornecedor) AS fornecedor, h.revenda, h.details,
                h.data_operacao AS data_emprestimo,
                (SELECT MIN(hc.data_operacao) FROM history hc WHERE hc.item_id=h.item_id AND hc.operation='Confirmação Empréstimo' AND hc.id>h.id AND hc.is_reversed=0) AS data_confirmacao,
                (SELECT MIN(hd.data_operacao) FROM history hd WHERE hd.item_id=h.item_id AND hd.operation='Devolução' AND hd.id>h.id AND hd.is_reversed=0) AS data_devolucao,
                'Empréstimo' AS operation_type,
                COALESCE(i.tipo, h.tipo) AS tipo, COALESCE(i.brand, h.brand) AS brand,
                COALESCE(i.model, h.model) AS model, COALESCE(i.identificador, h.identificador) AS identificador,
                COALESCE(i.nota_fiscal, h.nota_fiscal) AS nota_fiscal
            FROM history h LEFT JOIN items i ON i.id=h.item_id
            WHERE h.operation='Empréstimo' AND h.is_reversed=0 AND YEAR(h.data_operacao)=%s AND MONTH(h.data_operacao)=%s)
            UNION ALL
            (SELECT h.id, h.item_id, NULL, h.operador, NULL, NULL, NULL, NULL, NULL,
                i.fornecedor, i.revenda, h.details, h.data_operacao, NULL, NULL,
                'Cadastro', COALESCE(i.tipo,h.tipo), COALESCE(i.brand,h.brand),
                COALESCE(i.model,h.model), COALESCE(i.identificador,h.identificador), COALESCE(i.nota_fiscal,h.nota_fiscal)
            FROM history h LEFT JOIN items i ON i.id=h.item_id
            WHERE h.operation='Cadastro' AND h.is_reversed=0 AND YEAR(h.data_operacao)=%s AND MONTH(h.data_operacao)=%s)
            UNION ALL
            (SELECT h.id, h.item_id, h.peripheral_id, h.operador, NULL, NULL, NULL, NULL, NULL,
                NULL, NULL, h.details, h.data_operacao, NULL, NULL,
                h.operation, p.tipo, p.brand, p.model, p.identificador, NULL
            FROM history h LEFT JOIN peripherals p ON p.id=h.peripheral_id
            WHERE h.operation IN ('Cadastro Periférico','Vínculo Periférico','Desvínculo Periférico','Substituição Periférico')
            AND h.is_reversed=0 AND YEAR(h.data_operacao)=%s AND MONTH(h.data_operacao)=%s)
            UNION ALL
            (SELECT h.id, h.item_id, NULL, h.operador, NULL, NULL, NULL, NULL, NULL,
                h.fornecedor, h.revenda, h.details, h.data_operacao, NULL, NULL,
                'Exclusão', h.tipo, h.brand, h.model, h.identificador, h.nota_fiscal
            FROM history h
            WHERE h.operation='Exclusão' AND h.is_reversed=0 AND YEAR(h.data_operacao)=%s AND MONTH(h.data_operacao)=%s)
            ORDER BY data_emprestimo, item_id
        """
        with get_cursor(commit=False) as cur:
            cur.execute(sql, (ano, mes, ano, mes, ano, mes, ano, mes))
            return cur.fetchall()

    def reverse_history_entry(self, history_id: int, logged_user: str):
        try:
            with get_cursor() as cur:
                cur.execute("SELECT * FROM history WHERE id=%s", (history_id,))
                entry = cur.fetchone()
                if not entry:
                    return False, "Lançamento não encontrado."
                if entry["is_reversed"]:
                    return False, "Esta operação já foi estornada."

                op = entry["operation"]
                item_id = entry["item_id"]

                if op == "Confirmação Empréstimo":
                    cur.execute("UPDATE items SET status='Pendente' WHERE id=%s", (item_id,))
                elif op == "Empréstimo":
                    cur.execute("UPDATE items SET status='Disponível', assigned_to=NULL, cpf=NULL, date_issued=NULL, pessoa_juridica=0 WHERE id=%s", (item_id,))
                elif op == "Confirmação Devolução":
                    cur.execute("UPDATE items SET status='Pendente Devolução' WHERE id=%s", (item_id,))
                elif op == "Devolução":
                    # pessoa_juridica também é restaurado a partir do histórico do
                    # empréstimo original — confirm_return() zera esse campo em
                    # items, então sem essa leitura o estorno da devolução perderia
                    # a informação de que o empréstimo era de pessoa jurídica.
                    cur.execute("""
                        SELECT usuario, cpf, data_operacao, pessoa_juridica FROM history
                        WHERE item_id=%s AND operation IN ('Empréstimo','Confirmação Empréstimo') AND id<%s AND is_reversed=0
                        ORDER BY data_operacao DESC, id DESC LIMIT 1
                    """, (item_id, history_id))
                    last_loan = cur.fetchone()
                    if not last_loan:
                        return False, "Não foi possível encontrar o empréstimo original."
                    cur.execute(
                        "UPDATE items SET status='Indisponível', assigned_to=%s, cpf=%s, date_issued=%s, pessoa_juridica=%s WHERE id=%s",
                        (last_loan["usuario"], last_loan["cpf"], last_loan["data_operacao"], last_loan.get("pessoa_juridica") or 0, item_id),
                    )
                elif op == "Cadastro":
                    cur.execute("UPDATE items SET is_active=0 WHERE id=%s", (item_id,))
                else:
                    return False, f"Não é possível estornar uma operação do tipo '{op}'."

                cur.execute("UPDATE history SET is_reversed=1 WHERE id=%s", (history_id,))
                cur.execute(
                    """INSERT INTO history (item_id, operador, data_operacao, operation, usuario, cpf, cargo, center_cost, revenda)
                    VALUES (%s,%s,%s,'Estorno',%s,%s,%s,%s,%s)""",
                    (
                        item_id, logged_user, datetime.now(),
                        entry.get("usuario"), entry.get("cpf"),
                        entry.get("cargo"), entry.get("center_cost"), entry.get("revenda"),
                    ),
                )
            return True, f"Operação '{op}' do item {item_id} estornada com sucesso."
        except pymysql.MySQLError:
            logger.exception("Erro ao estornar")
            return False, "Erro ao estornar a operação."

    # ── Charts ─────────────────────────────────────────────────────────────────

    def get_issue_return_counts(self, year: int, month: int):
        num_days = calendar.monthrange(year, month)[1]
        days = list(range(1, num_days + 1))
        issues = {d: 0 for d in days}
        returns = {d: 0 for d in days}
        with get_cursor(commit=False) as cur:
            cur.execute(
                """SELECT DAY(data_operacao) as dia, operation, COUNT(id) as total
                FROM history
                WHERE YEAR(data_operacao)=%s AND MONTH(data_operacao)=%s
                AND operation IN ('Empréstimo','Devolução') AND is_reversed=0
                GROUP BY dia, operation""",
                (year, month),
            )
            rows = cur.fetchall()
        for row in rows:
            if row["operation"] == "Empréstimo":
                issues[row["dia"]] = row["total"]
            elif row["operation"] == "Devolução":
                returns[row["dia"]] = row["total"]
        return days, list(issues.values()), list(returns.values())

    def get_registration_counts(self, year: int, month: int):
        num_days = calendar.monthrange(year, month)[1]
        days = list(range(1, num_days + 1))
        registrations = {d: 0 for d in days}
        with get_cursor(commit=False) as cur:
            cur.execute(
                """SELECT DAY(data_operacao) as dia, COUNT(id) as total
                FROM history
                WHERE YEAR(data_operacao)=%s AND MONTH(data_operacao)=%s
                AND operation='Cadastro' AND is_reversed=0
                GROUP BY dia""",
                (year, month),
            )
            rows = cur.fetchall()
        for row in rows:
            registrations[row["dia"]] = row["total"]
        return days, list(registrations.values())
